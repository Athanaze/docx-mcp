"""
Formatting operations: format_text, format_table, table cell operations,
column widths, cell merging, custom styles.

All tools use unified block_index to address block items.
"""
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from word_document_server.document import docx_tool
from word_document_server.operations.helpers import parse_color, apply_run_format
from word_document_server.operations.blocks import (
    resolve_paragraph_block, resolve_table_block,
    get_block_items,
)


# ---------------------------------------------------------------------------
# Text formatting
# ---------------------------------------------------------------------------

def _capture_run_format(run):
    """Snapshot the explicit formatting of a run."""
    return {
        'bold': run.bold, 'italic': run.italic, 'underline': run.underline,
        'font_name': run.font.name, 'font_size': run.font.size,
        'color_rgb': run.font.color.rgb if run.font.color else None,
    }


def _restore_run_format(run, fmt):
    """Re-apply a captured format snapshot to a run."""
    if fmt.get('bold') is not None:
        run.bold = fmt['bold']
    if fmt.get('italic') is not None:
        run.italic = fmt['italic']
    if fmt.get('underline') is not None:
        run.underline = fmt['underline']
    if fmt.get('font_name'):
        run.font.name = fmt['font_name']
    if fmt.get('font_size'):
        run.font.size = fmt['font_size']
    if fmt.get('color_rgb'):
        run.font.color.rgb = fmt['color_rgb']


@docx_tool()
def format_text(doc, filename, block_index=None, start_pos=None, end_pos=None,
                search_text=None, match_occurrence=1,
                bold=None, italic=None, underline=None,
                color=None, font_size=None, font_name=None):
    """Format text by search_text match OR by block_index + character positions.

    Preferred: provide search_text (and optionally block_index to narrow scope).
    Legacy: provide block_index + start_pos + end_pos for exact character range.
    match_occurrence: which occurrence to format when search_text matches multiple
    times (1-based, default 1).
    """
    if search_text:
        items = get_block_items(doc)
        scope = [bi for bi in items if bi.type != "table"]
        if block_index is not None:
            try:
                bi = resolve_paragraph_block(doc, block_index)
                scope = [bi]
            except ValueError as e:
                return str(e)

        occ = 0
        for bi in scope:
            para = bi.obj
            text = para.text
            idx = -1
            search_from = 0
            while True:
                idx = text.find(search_text, search_from)
                if idx == -1:
                    break
                occ += 1
                if occ == int(match_occurrence):
                    break
                search_from = idx + 1
            if occ == int(match_occurrence) and idx != -1:
                start_pos = idx
                end_pos = idx + len(search_text)
                block_index = bi.index
                break

        if occ < int(match_occurrence):
            return f"Text '{search_text}' occurrence {match_occurrence} not found"

    if block_index is None or start_pos is None or end_pos is None:
        return "Must provide either search_text or (block_index + start_pos + end_pos)"

    try:
        bi = resolve_paragraph_block(doc, block_index)
    except ValueError as e:
        return str(e)

    start_pos = int(start_pos)
    end_pos = int(end_pos)
    para = bi.obj
    text = para.text
    if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
        return f"Invalid text positions. Paragraph has {len(text)} characters."

    target_text = text[start_pos:end_pos]

    runs = list(para.runs)
    if not runs:
        r = para.add_run(target_text)
        apply_run_format(r, bold=bold, italic=italic, underline=underline,
                          color=color, font_size=font_size, font_name=font_name)
        return f"Formatted text in block {block_index}: '{target_text}'"

    # Walk runs, find which ones overlap [start_pos, end_pos), split them,
    # and apply formatting only to the target range — preserving all other runs.
    new_runs = []
    char_offset = 0
    for run in runs:
        run_text = run.text or ''
        run_start = char_offset
        run_end = char_offset + len(run_text)
        fmt = _capture_run_format(run)

        if run_end <= start_pos or run_start >= end_pos:
            new_runs.append((run_text, fmt, False))
        elif run_start >= start_pos and run_end <= end_pos:
            new_runs.append((run_text, fmt, True))
        else:
            if run_start < start_pos:
                new_runs.append((run_text[:start_pos - run_start], fmt, False))
            overlap_s = max(start_pos, run_start) - run_start
            overlap_e = min(end_pos, run_end) - run_start
            new_runs.append((run_text[overlap_s:overlap_e], fmt, True))
            if run_end > end_pos:
                new_runs.append((run_text[end_pos - run_start:], fmt, False))

        char_offset = run_end

    for run in runs:
        run._element.getparent().remove(run._element)

    for text_part, fmt, is_target in new_runs:
        if not text_part:
            continue
        r = para.add_run(text_part)
        _restore_run_format(r, fmt)
        if is_target:
            apply_run_format(r, bold=bold, italic=italic, underline=underline,
                              color=color, font_size=font_size, font_name=font_name)

    return f"Formatted text in block {block_index}: '{target_text}'"


# ---------------------------------------------------------------------------
# Custom styles
# ---------------------------------------------------------------------------

@docx_tool()
def create_style(doc, filename, style_name, bold=None, italic=None,
                 font_size=None, font_name=None, color=None, base_style=None):
    try:
        style = doc.styles[style_name]
        return f"Style '{style_name}' already exists"
    except KeyError:
        pass

    new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    if base_style:
        try:
            new_style.base_style = doc.styles[base_style]
        except KeyError:
            pass

    font = new_style.font
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if font_size is not None:
        font.size = Pt(int(font_size))
    if font_name:
        font.name = font_name
    if color:
        rgb = parse_color(color)
        if rgb:
            font.color.rgb = rgb

    return f"Style '{style_name}' created in {filename}"


# ---------------------------------------------------------------------------
# Table formatting (addressed by block_index)
# ---------------------------------------------------------------------------

def _get_table(doc, block_index):
    """Resolve block_index to a Table, returning (table, error_string)."""
    try:
        bi = resolve_table_block(doc, block_index)
        return bi.obj, None
    except ValueError as e:
        return None, str(e)


def _set_cell_shading(cell, fill_color, pattern="clear"):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="{pattern}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


@docx_tool()
def format_table(doc, filename, block_index, has_header_row=None,
                 border_style=None):
    table, err = _get_table(doc, block_index)
    if err:
        return err

    if has_header_row:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

    if border_style:
        val = border_style if border_style != 'none' else 'nil'
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')
                for edge in ['top', 'left', 'bottom', 'right']:
                    el = OxmlElement(f'w:{edge}')
                    el.set(qn('w:val'), val)
                    el.set(qn('w:sz'), '4')
                    el.set(qn('w:space'), '0')
                    el.set(qn('w:color'), 'auto')
                    borders.append(el)
                tcPr.append(borders)

    return f"Table at block {block_index} formatted in {filename}"


@docx_tool()
def set_table_cell_shading(doc, filename, block_index, row_index, col_index,
                           fill_color, pattern="clear"):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    row_index, col_index = int(row_index), int(col_index)
    if row_index >= len(table.rows) or col_index >= len(table.columns):
        return "Cell index out of range"
    _set_cell_shading(table.cell(row_index, col_index), fill_color, pattern)
    return f"Shading applied to cell ({row_index},{col_index}) in table at block {block_index}"


@docx_tool()
def apply_table_alternating_rows(doc, filename, block_index,
                                 color1="FFFFFF", color2="F2F2F2"):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    for ri, row in enumerate(table.rows):
        color = color1 if ri % 2 == 0 else color2
        for cell in row.cells:
            _set_cell_shading(cell, color)
    return f"Alternating row colors applied to table at block {block_index}"


@docx_tool()
def highlight_table_header(doc, filename, block_index,
                           header_color="4472C4", text_color="FFFFFF"):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, header_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                rgb = parse_color(text_color)
                if rgb:
                    run.font.color.rgb = rgb
    return f"Header highlighted in table at block {block_index}"


@docx_tool()
def merge_cells(doc, filename, block_index, start_row, start_col, end_row, end_col):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    start_row, start_col = int(start_row), int(start_col)
    end_row, end_col = int(end_row), int(end_col)
    a = table.cell(start_row, start_col)
    b = table.cell(end_row, end_col)
    a.merge(b)
    return f"Cells ({start_row},{start_col})-({end_row},{end_col}) merged in table at block {block_index}"


@docx_tool()
def merge_cells_horizontal(doc, filename, block_index, row, start_col, end_col):
    """Merge cells horizontally in one row."""
    table, err = _get_table(doc, block_index)
    if err:
        return err
    row, start_col, end_col = int(row), int(start_col), int(end_col)
    table.cell(row, start_col).merge(table.cell(row, end_col))
    return f"Cells ({row},{start_col})-({row},{end_col}) merged in table at block {block_index}"


@docx_tool()
def merge_cells_vertical(doc, filename, block_index, col, start_row, end_row):
    """Merge cells vertically in one column."""
    table, err = _get_table(doc, block_index)
    if err:
        return err
    col, start_row, end_row = int(col), int(start_row), int(end_row)
    table.cell(start_row, col).merge(table.cell(end_row, col))
    return f"Cells ({start_row},{col})-({end_row},{col}) merged in table at block {block_index}"


def _convert_width(width, width_type):
    """Convert width from user units to appropriate internal value."""
    width = float(width)
    if width_type == "inches":
        return int(width * 914400), "dxa"
    elif width_type == "cm":
        return int(width * 360000 / 914400 * 1440), "dxa"
    elif width_type == "percent":
        return int(width * 50), "pct"
    else:
        return int(width * 20), "dxa"


@docx_tool()
def set_column_widths(doc, filename, block_index, widths, width_type="points"):
    table, err = _get_table(doc, block_index)
    if err:
        return err

    for ci, w in enumerate(widths):
        if ci >= len(table.columns):
            break
        tw, tp = _convert_width(w, width_type)
        for row in table.rows:
            tc = row.cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(tw))
            tcW.set(qn('w:type'), tp)

    return f"Column widths set for table at block {block_index}"


@docx_tool()
def set_cell_alignment(doc, filename, block_index, row_index, col_index,
                       horizontal="left", vertical="top"):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    cell = table.cell(int(row_index), int(col_index))
    h_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
             "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    v_map = {"top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
             "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
             "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM}
    for para in cell.paragraphs:
        para.alignment = h_map.get(horizontal, WD_ALIGN_PARAGRAPH.LEFT)
    cell.vertical_alignment = v_map.get(vertical, WD_CELL_VERTICAL_ALIGNMENT.TOP)
    return f"Alignment set for cell ({row_index},{col_index}) in table at block {block_index}"


@docx_tool()
def format_cell_text(doc, filename, block_index, row_index, col_index,
                     text_content=None, bold=None, italic=None, underline=None,
                     color=None, font_size=None, font_name=None):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    cell = table.cell(int(row_index), int(col_index))
    if text_content is not None:
        cell.text = str(text_content)
    for para in cell.paragraphs:
        for run in para.runs:
            apply_run_format(run, bold=bold, italic=italic, underline=underline,
                              color=color, font_size=font_size, font_name=font_name)
    return f"Cell text formatted at ({row_index},{col_index}) in table at block {block_index}"


@docx_tool()
def set_cell_padding(doc, filename, block_index, row_index, col_index,
                     top=None, bottom=None, left=None, right=None, unit="points"):
    table, err = _get_table(doc, block_index)
    if err:
        return err
    cell = table.cell(int(row_index), int(col_index))
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)

    def _to_twips(val):
        v = float(val)
        if unit == "inches":
            return int(v * 1440)
        elif unit == "cm":
            return int(v * 567)
        return int(v * 20)

    for edge, value in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        if value is not None:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:w'), str(_to_twips(value)))
            el.set(qn('w:type'), 'dxa')
            existing = tcMar.find(qn(f'w:{edge}'))
            if existing is not None:
                tcMar.remove(existing)
            tcMar.append(el)

    return f"Padding set for cell ({row_index},{col_index}) in table at block {block_index}"


@docx_tool()
def set_table_width(doc, filename, block_index, width, width_type="points"):
    """Set the overall width of a table.

    width_type: 'points', 'inches', 'cm', 'percent', or 'auto'.
    """
    table, err = _get_table(doc, block_index)
    if err:
        return err

    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn('w:tblW'))
    if existing is not None:
        tbl_pr.remove(existing)

    tbl_w = OxmlElement('w:tblW')
    if width_type == "auto":
        tbl_w.set(qn('w:w'), '0')
        tbl_w.set(qn('w:type'), 'auto')
    else:
        tw, tp = _convert_width(width, width_type)
        tbl_w.set(qn('w:w'), str(tw))
        tbl_w.set(qn('w:type'), tp)
    tbl_pr.append(tbl_w)

    return f"Table width set for table at block {block_index}"


@docx_tool()
def auto_fit_table(doc, filename, block_index):
    """Set table to auto-fit columns based on content."""
    table, err = _get_table(doc, block_index)
    if err:
        return err

    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)

    existing = tbl_pr.find(qn('w:tblLayout'))
    if existing is not None:
        tbl_pr.remove(existing)

    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'autofit')
    tbl_pr.append(layout)

    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), '0')
            tcW.set(qn('w:type'), 'auto')

    return f"Auto-fit applied to table at block {block_index}"


@docx_tool()
def set_table_alignment_all(doc, filename, block_index,
                            horizontal="left", vertical="top"):
    """Set text alignment for ALL cells in a table at once."""
    table, err = _get_table(doc, block_index)
    if err:
        return err

    h_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
             "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
    v_map = {"top": WD_CELL_VERTICAL_ALIGNMENT.TOP,
             "center": WD_CELL_VERTICAL_ALIGNMENT.CENTER,
             "bottom": WD_CELL_VERTICAL_ALIGNMENT.BOTTOM}

    h_align = h_map.get(horizontal, WD_ALIGN_PARAGRAPH.LEFT)
    v_align = v_map.get(vertical, WD_CELL_VERTICAL_ALIGNMENT.TOP)

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = h_align
            cell.vertical_alignment = v_align

    return f"Alignment set for all cells in table at block {block_index}"
