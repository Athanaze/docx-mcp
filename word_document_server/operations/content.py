"""
Content operations: create, read, update, delete for block items (paragraphs,
tables, rows, columns, lists). Also search/replace and block replacement.

All tools use a unified block_index that numbers every block-level item
(paragraph, heading, list item, table) in document order.
"""
import copy
import os
import json
import zipfile

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

from word_document_server.document import docx_tool
from word_document_server.paths import resolve_docx, resolve_path
from word_document_server.operations.numbering import (
    set_paragraph_list, create_restart_num_id,
)
from word_document_server.operations.blocks import (
    get_block_items, resolve_block, resolve_paragraph_block,
    resolve_table_block, find_block, normalize_text,
)
from word_document_server.operations.helpers import (
    parse_color, apply_run_format, resolve_list_style,
)

# Backwards-compatible aliases for imports from formatting.py
_parse_color = parse_color
_apply_run_format = apply_run_format


def _copy_run_formatting(src, dst):
    """Copy explicit run-level formatting from src to dst."""
    if src.font.name:
        dst.font.name = src.font.name
    if src.font.size:
        dst.font.size = src.font.size
    if src.bold is not None:
        dst.bold = src.bold
    if src.italic is not None:
        dst.italic = src.italic
    if src.underline is not None:
        dst.underline = src.underline
    if src.font.color and src.font.color.rgb:
        dst.font.color.rgb = src.font.color.rgb


def _xml_element(block_item_obj):
    """Get the underlying lxml element from a python-docx Paragraph or Table."""
    if isinstance(block_item_obj, Table):
        return block_item_obj._tbl
    return block_item_obj._element


# ---------------------------------------------------------------------------
# MCP Tools — Document lifecycle
# ---------------------------------------------------------------------------

@docx_tool(creates=True)
def create_document(doc, filename, title=None, author=None):
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author
    return f"Document created: {filename}"


@docx_tool(readonly=True)
def get_document_info(doc, filename):
    props = doc.core_properties
    items = get_block_items(doc)
    return json.dumps({
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "created": str(props.created) if props.created else "",
        "modified": str(props.modified) if props.modified else "",
        "last_modified_by": props.last_modified_by or "",
        "revision": props.revision or 0,
        "section_count": len(doc.sections),
        "word_count": sum(len(p.text.split()) for p in doc.paragraphs),
        "block_count": len(items),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }, indent=2)


@docx_tool(readonly=True)
def get_document_text(doc, filename, include_indices=True):
    """Extract text from a Word document in document order.

    Every block item (paragraph, heading, list item, table) gets a unified
    block_index shown as [N]. Use this index with all tools (format_text,
    delete_block, insert_content, add_table_row, etc.).
    """
    items = get_block_items(doc)
    lines = []
    for bi in items:
        if bi.type == "table":
            tbl = bi.obj
            if include_indices:
                lines.append(f"[{bi.index}] (Table {len(tbl.rows)}x{len(tbl.columns)})")
            for ri, row in enumerate(tbl.rows):
                cells = [row.cells[ci].text for ci in range(len(tbl.columns))]
                if include_indices:
                    lines.append(f"    Row {ri}: " + " | ".join(cells))
                else:
                    lines.append(" | ".join(cells))
        else:
            p = bi.obj
            if include_indices:
                style_hint = ""
                if bi.type == "heading":
                    style_hint = f" ({p.style.name})"
                elif bi.type == "list_item":
                    style_hint = f" ({p.style.name})"
                lines.append(f"[{bi.index}]{style_hint} {p.text}")
            else:
                lines.append(p.text)
    return "\n".join(lines)


@docx_tool(readonly=True)
def get_document_outline(doc, filename):
    items = get_block_items(doc)
    blocks = []
    for bi in items:
        if bi.type == "table":
            tbl = bi.obj
            preview = []
            for ri in range(min(3, len(tbl.rows))):
                row_data = []
                for ci in range(min(3, len(tbl.columns))):
                    try:
                        t = tbl.cell(ri, ci).text
                        row_data.append(t[:20] + ("..." if len(t) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                preview.append(row_data)
            blocks.append({
                "block_index": bi.index, "type": "table",
                "rows": len(tbl.rows), "columns": len(tbl.columns),
                "preview": preview,
            })
        else:
            blocks.append({
                "block_index": bi.index, "type": bi.type,
                "text": bi.obj.text[:100] + ("..." if len(bi.obj.text) > 100 else ""),
                "style": bi.obj.style.name if bi.obj.style else "Normal",
            })
    return json.dumps({"blocks": blocks}, indent=2)


def _run_to_dict(run):
    """Serialize a run for get_blocks JSON."""
    d = {"text": run.text or ""}
    if run.bold is not None:
        d["bold"] = run.bold
    if run.italic is not None:
        d["italic"] = run.italic
    if run.underline is not None:
        try:
            d["underline"] = bool(run.underline)
        except Exception:
            d["underline"] = True
    if run.font.size:
        d["font_size_pt"] = float(run.font.size.pt)
    if run.font.name:
        d["font_name"] = run.font.name
    if run.font.color and run.font.color.rgb:
        rgb = run.font.color.rgb
        d["color_hex"] = f"{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
    return d


@docx_tool(readonly=True)
def get_blocks(doc, filename, start_block_index=0, end_block_index=None,
               include_runs=True):
    """Return structured JSON for blocks in range (inclusive block_index).

    For non-table blocks includes style name and optional per-run formatting.
    For tables includes row/column counts and a 2D array of cell texts.
    """
    items = get_block_items(doc)
    n = len(items)
    start = max(0, int(start_block_index))
    end = n - 1 if end_block_index is None else int(end_block_index)
    if end < start:
        return json.dumps({"blocks": [], "error": "end_block_index < start_block_index"}, indent=2)

    out = []
    for bi in items:
        if bi.index < start or bi.index > end:
            continue
        if bi.type == "table":
            tbl = bi.obj
            cell_texts = []
            for row in tbl.rows:
                cell_texts.append([c.text for c in row.cells])
            out.append({
                "block_index": bi.index,
                "type": "table",
                "rows": len(tbl.rows),
                "columns": len(tbl.columns),
                "cell_texts": cell_texts,
            })
        else:
            para = bi.obj
            entry = {
                "block_index": bi.index,
                "type": bi.type,
                "text": para.text,
                "style": para.style.name if para.style else None,
            }
            if include_runs:
                entry["runs"] = [_run_to_dict(r) for r in para.runs]
            out.append(entry)
    return json.dumps({"blocks": out, "block_count": n}, indent=2)


@docx_tool(readonly=True)
def list_document_styles(doc, filename):
    """List paragraph and character styles available in the document."""
    rows = []
    for style in doc.styles:
        try:
            rows.append({
                "name": style.name,
                "type": style.type.name if style.type else None,
            })
        except Exception:
            continue
    rows.sort(key=lambda x: (x["name"] or "").lower())
    return json.dumps({"styles": rows, "count": len(rows)}, indent=2)


@docx_tool()
def set_paragraph_style(doc, filename, block_index, style_name):
    """Apply a document style to the paragraph at the given block_index."""
    bi = resolve_paragraph_block(doc, block_index)
    try:
        bi.obj.style = doc.styles[style_name]
    except KeyError:
        return f"Style '{style_name}' not found. Use list_document_styles to see available names."
    return f"Style '{style_name}' applied to block {block_index}"


def list_documents(directory="."):
    from word_document_server.paths import resolve_directory
    try:
        resolved = resolve_directory(directory)
    except ValueError as e:
        return str(e)
    if not os.path.isdir(resolved):
        return f"Directory {resolved} does not exist"
    docs = [f for f in os.listdir(resolved) if f.endswith('.docx') and not f.startswith('~$')]
    return json.dumps({"directory": resolved, "documents": sorted(docs)}, indent=2)


def copy_document(source_filename, destination_filename=None):
    from word_document_server.paths import copy_document as _copy
    try:
        src = resolve_docx(source_filename)
    except ValueError as e:
        return str(e)
    dest = None
    if destination_filename:
        try:
            dest = resolve_docx(destination_filename)
        except ValueError as e:
            return str(e)
    ok, msg, path = _copy(src, dest)
    return msg


def merge_documents(target_filename, source_filenames, add_page_breaks=True):
    try:
        target = resolve_docx(target_filename)
    except ValueError as e:
        return str(e)
    target_doc = Document()
    first = True
    for sf in source_filenames:
        try:
            src_path = resolve_docx(sf)
        except ValueError as e:
            return str(e)
        if not os.path.exists(src_path):
            return f"Source document {src_path} does not exist"
        src_doc = Document(src_path)
        if not first and add_page_breaks:
            target_doc.add_page_break()

        for item in src_doc.iter_inner_content():
            if isinstance(item, Paragraph):
                new_para = target_doc.add_paragraph(style=item.style)
                for run in item.runs:
                    new_run = new_para.add_run(run.text)
                    _copy_run_formatting(run, new_run)
                pPr = item._element.find(qn('w:pPr'))
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        new_pPr = new_para._element.get_or_add_pPr()
                        existing = new_pPr.find(qn('w:numPr'))
                        if existing is not None:
                            new_pPr.remove(existing)
                        new_pPr.append(copy.deepcopy(numPr))
            elif isinstance(item, Table):
                rows = len(item.rows)
                cols = len(item.columns)
                new_table = target_doc.add_table(rows=rows, cols=cols)
                try:
                    new_table.style = item.style
                except Exception:
                    pass
                for ri, row in enumerate(item.rows):
                    for ci, cell in enumerate(row.cells):
                        new_table.cell(ri, ci).text = cell.text
        first = False
    target_doc.save(target)
    return f"Merged {len(source_filenames)} documents into {target}"


def get_document_xml(filename):
    try:
        path = resolve_docx(filename)
    except ValueError as e:
        return str(e)
    if not os.path.exists(path):
        return f"Document {path} does not exist"
    try:
        with zipfile.ZipFile(path) as z:
            with z.open('word/document.xml') as f:
                return f.read().decode('utf-8')
    except Exception as e:
        return f"Failed to extract XML: {e}"


# ---------------------------------------------------------------------------
# Content addition tools
# ---------------------------------------------------------------------------

@docx_tool()
def add_table_of_contents(doc, filename, title="Table of Contents",
                          max_level=3):
    """Insert a TOC field code that Word/LibreOffice will update on open.

    This inserts a proper w:fldSimple TOC field. The TOC is rendered when the
    document is opened in Word or LibreOffice (or when fields are updated).
    """
    max_level = max(1, min(9, int(max_level)))

    toc_heading = doc.add_paragraph(title, style='Heading 1')

    toc_para = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f' TOC \\o "1-{max_level}" \\h \\z \\u ')
    run = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "Right-click and select 'Update Field' to generate TOC"
    run.append(t)
    fld.append(run)
    toc_para._element.append(fld)

    return f"Table of Contents (levels 1-{max_level}) added to {filename}"


@docx_tool()
def add_heading(doc, filename, text, level=1, font_name=None, font_size=None,
                bold=None, italic=None, border_bottom=False):
    level = max(1, min(9, int(level)))
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _apply_run_format(run, bold=bold, italic=italic, font_size=font_size,
                          font_name=font_name)
    if border_bottom:
        pPr = heading._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return f"Heading (level {level}) added to {filename}"


@docx_tool()
def add_paragraph(doc, filename, text, style=None, font_name=None,
                  font_size=None, bold=None, italic=None, color=None):
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        _apply_run_format(run, bold=bold, italic=italic, font_name=font_name,
                          font_size=font_size, color=color)
    return f"Paragraph added to {filename}"


@docx_tool()
def add_table(doc, filename, rows, cols, data=None):
    rows = int(rows)
    cols = int(cols)
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = 'Table Grid'
    except Exception:
        pass
    if data:
        for ri, row_data in enumerate(data):
            if ri >= rows:
                break
            for ci, cell_text in enumerate(row_data):
                if ci >= cols:
                    break
                table.cell(ri, ci).text = str(cell_text)
    return f"Table ({rows}x{cols}) added to {filename}"


@docx_tool()
def add_picture(doc, filename, image_path, width=None):
    try:
        resolved_image = resolve_path(image_path)
    except ValueError as e:
        return str(e)
    if not os.path.exists(resolved_image):
        return f"Image file {resolved_image} does not exist"
    kwargs = {}
    if width:
        kwargs['width'] = Inches(float(width))
    doc.add_picture(resolved_image, **kwargs)
    return f"Picture added to {filename}"


@docx_tool()
def add_page_break(doc, filename):
    doc.add_page_break()
    return f"Page break added to {filename}"


@docx_tool()
def add_list(doc, filename, items, list_type="bullet", level=0):
    """Add a bulleted or numbered list with proper numbering definitions."""
    if not items:
        return "No items provided"
    num_id = create_restart_num_id(doc, list_type)

    style_name = resolve_list_style(doc)

    for item in items:
        para = doc.add_paragraph(str(item), style=style_name)
        set_paragraph_list(para, num_id, level=int(level))

    kind = "bulleted" if list_type == "bullet" else "numbered"
    return f"{kind.capitalize()} list with {len(items)} items added to {filename}"


# ---------------------------------------------------------------------------
# Insert content near existing block item
# ---------------------------------------------------------------------------

@docx_tool()
def insert_content(doc, filename, content_type="paragraph", text="",
                   target_text=None, target_block_index=None,
                   position="after", style=None,
                   items=None, list_type="bullet", level=1,
                   table_rows=None, table_cols=None, table_data=None):
    """
    Insert a heading, paragraph, list, or table before/after a target block
    identified by text or block_index.
    """
    bi = find_block(doc, target_text=target_text, block_index=target_block_index)
    if bi is None:
        return "Target block not found (TOC paragraphs are skipped in text search)"

    target_el = _xml_element(bi.obj)

    if content_type == "heading":
        lv = max(1, min(9, int(level)))
        new_el = doc.add_paragraph(text, style=style or f'Heading {lv}')._element
    elif content_type == "table":
        rows = int(table_rows or 2)
        cols = int(table_cols or 2)
        tbl = doc.add_table(rows=rows, cols=cols)
        try:
            tbl.style = 'Table Grid'
        except Exception:
            pass
        if table_data:
            for ri, row_data in enumerate(table_data):
                if ri >= rows:
                    break
                for ci, cell_text in enumerate(row_data):
                    if ci >= cols:
                        break
                    tbl.cell(ri, ci).text = str(cell_text)
        new_el = tbl._tbl
    elif content_type == "list":
        if not items:
            return "No list items provided"
        num_id = create_restart_num_id(doc, list_type)
        list_style = resolve_list_style(doc)
        new_paras = []
        for item in items:
            p = doc.add_paragraph(str(item), style=list_style)
            set_paragraph_list(p, num_id, level=0)
            new_paras.append(p)
        for p in reversed(new_paras):
            if position == 'before':
                target_el.addprevious(p._element)
            else:
                target_el.addnext(p._element)
        kind = "bulleted" if list_type == "bullet" else "numbered"
        return f"{kind.capitalize()} list with {len(items)} items inserted {position} block {bi.index}"
    else:
        if bi.type != "table":
            para = bi.obj
            if style:
                use_style = style
            elif para.style and para.style.name.lower().startswith('heading'):
                use_style = 'Normal'
            else:
                use_style = para.style
            new_p = doc.add_paragraph(text, style=use_style)
            if not style and para.runs and new_p.runs:
                if not (para.style and para.style.name.lower().startswith('heading')):
                    _copy_run_formatting(para.runs[0], new_p.runs[0])
        else:
            new_p = doc.add_paragraph(text, style=style or 'Normal')
        new_el = new_p._element

    if position == 'before':
        target_el.addprevious(new_el)
    else:
        target_el.addnext(new_el)

    return f"{content_type.capitalize()} inserted {position} block {bi.index}"


# ---------------------------------------------------------------------------
# Delete, move, search/replace
# ---------------------------------------------------------------------------

@docx_tool()
def delete_block(doc, filename, block_index):
    """Delete any block item (paragraph, heading, list item, or table) by its index."""
    try:
        bi = resolve_block(doc, block_index)
    except ValueError as e:
        return str(e)
    el = _xml_element(bi.obj)
    el.getparent().remove(el)
    return f"Block {block_index} ({bi.type}) deleted from {filename}"


@docx_tool()
def move_block(doc, filename, source_index, target_index, position="after"):
    """Move any block item from source_index to before/after target_index."""
    try:
        src_bi = resolve_block(doc, source_index)
    except ValueError as e:
        return str(e)
    try:
        tgt_bi = resolve_block(doc, target_index)
    except ValueError as e:
        return str(e)
    if source_index == target_index:
        return "Source and target are the same block"

    src_el = _xml_element(src_bi.obj)
    tgt_el = _xml_element(tgt_bi.obj)

    src_el.getparent().remove(src_el)
    if position == "before":
        tgt_el.addprevious(src_el)
    else:
        tgt_el.addnext(src_el)

    return f"Block {source_index} ({src_bi.type}) moved {position} block {target_index}"


@docx_tool()
def search_and_replace(doc, filename, find_text, replace_text):
    count = 0
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("TOC"):
            continue
        if find_text in para.text:
            count += _replace_in_runs(para, find_text, replace_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.style and para.style.name.startswith("TOC"):
                        continue
                    if find_text in para.text:
                        count += _replace_in_runs(para, find_text, replace_text)
    if count == 0:
        return f"Text '{find_text}' not found in {filename}"
    return f"Replaced {count} occurrence(s) of '{find_text}' in {filename}"


def _replace_in_runs(paragraph, old_text, new_text):
    """Cross-run text replacement preserving formatting."""
    runs = paragraph.runs
    if not runs:
        return 0
    count = 0
    while True:
        texts = [r.text or '' for r in runs]
        full = ''.join(texts)
        start = full.find(old_text)
        if start == -1:
            break
        end = start + len(old_text)
        offsets = []
        off = 0
        for t in texts:
            offsets.append(off)
            off += len(t)
        first_ri = last_ri = None
        for i, s in enumerate(offsets):
            e = s + len(texts[i])
            if first_ri is None and e > start:
                first_ri = i
            if s < end:
                last_ri = i
        if first_ri is None or last_ri is None:
            break
        if first_ri == last_ri:
            runs[first_ri].text = runs[first_ri].text.replace(old_text, new_text, 1)
        else:
            off_in_first = start - offsets[first_ri]
            runs[first_ri].text = runs[first_ri].text[:off_in_first] + new_text
            for i in range(first_ri + 1, last_ri):
                runs[i].text = ''
            off_in_last = end - offsets[last_ri]
            runs[last_ri].text = runs[last_ri].text[off_in_last:]
        count += 1
        runs = paragraph.runs
    return count


@docx_tool(readonly=True)
def find_text(doc, filename, text_to_find, match_case=True, whole_word=False):
    """Find text in paragraphs and table cells. Returns block_index for each match."""
    import re as _re
    results = []
    if not text_to_find:
        return json.dumps({"matches": results, "count": 0}, indent=2)
    search = text_to_find if match_case else text_to_find.lower()
    items = get_block_items(doc)

    def _matches(text):
        t = text if match_case else text.lower()
        if whole_word:
            return bool(_re.search(r'\b' + _re.escape(search) + r'\b', t))
        return search in t

    for bi in items:
        if bi.type == "table":
            for ri, row in enumerate(bi.obj.rows):
                for ci, cell in enumerate(row.cells):
                    if _matches(cell.text):
                        results.append({
                            "block_index": bi.index, "type": "table_cell",
                            "row": ri, "col": ci,
                            "text": cell.text[:100],
                        })
        else:
            if _matches(bi.obj.text):
                results.append({"block_index": bi.index, "type": bi.type,
                                "text": bi.obj.text[:100]})
    return json.dumps({"matches": results, "count": len(results)}, indent=2)


# ---------------------------------------------------------------------------
# replace_block — operates on raw XML elements
# ---------------------------------------------------------------------------

def _extract_element_text(el):
    return "".join(n.text or '' for n in el.iter() if n.tag.endswith('}t'))


def _get_element_style(el):
    pPr = el.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            return pStyle.get(qn('w:val'))
    return None


_W_P = qn('w:p')
_W_TBL = qn('w:tbl')


@docx_tool()
def replace_block(doc, filename, header_text=None, start_anchor=None,
                  end_anchor=None, new_paragraphs=None, style=None):
    """
    Replace a block of content. Two modes:
    - header_text: replaces everything below that heading until the next heading/TOC.
    - start_anchor/end_anchor: replaces everything between those text markers.
    """
    if not new_paragraphs:
        new_paragraphs = []

    body = doc.element.body
    elements = list(body)

    if header_text:
        target_norm = normalize_text(header_text)
        header_el = None
        header_idx = None
        for i, el in enumerate(elements):
            if el.tag != _W_P:
                continue
            sn = _get_element_style(el)
            if sn and sn.upper().startswith("TOC"):
                continue
            if normalize_text(_extract_element_text(el)) == target_norm:
                header_el, header_idx = el, i
                break
        if header_el is None:
            for i, el in enumerate(elements):
                if el.tag != _W_P:
                    continue
                sn = _get_element_style(el)
                if sn and sn.upper().startswith("TOC"):
                    continue
                en = normalize_text(_extract_element_text(el))
                if target_norm in en or en in target_norm:
                    header_el, header_idx = el, i
                    break
        if header_el is None:
            return f"Header '{header_text}' not found"

        to_remove = []
        for i in range(header_idx + 1, len(elements)):
            el = elements[i]
            if el.tag == _W_P:
                sn = _get_element_style(el)
                if sn and sn.lower().startswith(('heading', 'toc')):
                    break
            to_remove.append(el)
        for el in to_remove:
            body.remove(el)

        cur = header_el
        for text in new_paragraphs:
            p = doc.add_paragraph(text, style=style or "Normal")
            cur.addnext(p._element)
            cur = p._element

        return (f"Replaced content under '{header_text}' with "
                f"{len(new_paragraphs)} paragraph(s), removed {len(to_remove)} elements.")

    elif start_anchor:
        start_norm = normalize_text(start_anchor)
        start_idx = None
        for i, el in enumerate(elements):
            if el.tag != _W_P:
                continue
            if normalize_text(_extract_element_text(el)) == start_norm:
                start_idx = i
                break
        if start_idx is None:
            for i, el in enumerate(elements):
                if el.tag != _W_P:
                    continue
                en = normalize_text(_extract_element_text(el))
                if start_norm in en or en in start_norm:
                    start_idx = i
                    break
        if start_idx is None:
            return f"Start anchor '{start_anchor}' not found"

        end_idx = None
        if end_anchor:
            end_norm = normalize_text(end_anchor)
            for i in range(start_idx + 1, len(elements)):
                if elements[i].tag != _W_P:
                    continue
                if normalize_text(_extract_element_text(elements[i])) == end_norm:
                    end_idx = i
                    break
            if end_idx is None:
                for i in range(start_idx + 1, len(elements)):
                    if elements[i].tag != _W_P:
                        continue
                    en = normalize_text(_extract_element_text(elements[i]))
                    if end_norm in en or en in end_norm:
                        end_idx = i
                        break
        else:
            for i in range(start_idx + 1, len(elements)):
                if elements[i].tag != _W_P:
                    continue
                sn = _get_element_style(elements[i])
                if sn and sn.lower().startswith(('heading', 'toc')):
                    end_idx = i
                    break

        to_remove = []
        for i in range(start_idx + 1, end_idx if end_idx is not None else len(elements)):
            to_remove.append(elements[i])
        for el in to_remove:
            body.remove(el)

        cur = elements[start_idx]
        for text in new_paragraphs:
            p = doc.add_paragraph(text, style=style or "Normal")
            cur.addnext(p._element)
            cur = p._element

        return (f"Replaced content between '{start_anchor}' and "
                f"'{end_anchor or 'next heading'}' with {len(new_paragraphs)} "
                f"paragraph(s), removed {len(to_remove)} elements.")

    return "Must specify either header_text or start_anchor"


# ---------------------------------------------------------------------------
# Table structure operations (addressed by block_index)
# ---------------------------------------------------------------------------

@docx_tool()
def add_table_row(doc, filename, block_index, row_data=None, position="end"):
    """Add a row to an existing table identified by block_index."""
    try:
        bi = resolve_table_block(doc, block_index)
    except ValueError as e:
        return str(e)
    table = bi.obj

    cols = len(table.columns)
    new_tr = OxmlElement('w:tr')
    for ci in range(cols):
        tc = OxmlElement('w:tc')
        p = OxmlElement('w:p')
        if row_data and ci < len(row_data):
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = str(row_data[ci])
            r.append(t)
            p.append(r)
        tc.append(p)
        new_tr.append(tc)

    tbl = table._tbl
    if position == "start":
        first_tr = tbl.find(qn('w:tr'))
        if first_tr is not None:
            first_tr.addprevious(new_tr)
        else:
            tbl.append(new_tr)
    elif position == "end":
        tbl.append(new_tr)
    else:
        idx = int(position)
        rows = tbl.findall(qn('w:tr'))
        if idx < 0 or idx >= len(rows):
            tbl.append(new_tr)
        else:
            rows[idx].addprevious(new_tr)

    return f"Row added to table at block {block_index}, position '{position}'"


@docx_tool()
def delete_table_row(doc, filename, block_index, row_index):
    """Delete a specific row from a table identified by block_index."""
    try:
        bi = resolve_table_block(doc, block_index)
    except ValueError as e:
        return str(e)
    table = bi.obj
    row_index = int(row_index)
    if row_index < 0 or row_index >= len(table.rows):
        return f"Invalid row index. Table has {len(table.rows)} rows."
    tr = table.rows[row_index]._tr
    tr.getparent().remove(tr)
    return f"Row {row_index} deleted from table at block {block_index}"


@docx_tool()
def add_table_column(doc, filename, block_index, col_data=None, position="end"):
    """Add a column to an existing table identified by block_index."""
    try:
        bi = resolve_table_block(doc, block_index)
    except ValueError as e:
        return str(e)
    table = bi.obj

    num_rows = len(table.rows)
    if not col_data:
        col_data = [""] * num_rows

    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        new_gc = OxmlElement('w:gridCol')
        existing_gcs = grid.findall(qn('w:gridCol'))
        if existing_gcs:
            w = existing_gcs[0].get(qn('w:w'))
            if w:
                new_gc.set(qn('w:w'), w)
        if position == "start":
            if existing_gcs:
                existing_gcs[0].addprevious(new_gc)
            else:
                grid.append(new_gc)
        elif position == "end":
            grid.append(new_gc)
        else:
            col_idx = int(position)
            if col_idx < len(existing_gcs):
                existing_gcs[col_idx].addprevious(new_gc)
            else:
                grid.append(new_gc)

    for ri, row in enumerate(table.rows):
        tc = OxmlElement('w:tc')
        p = OxmlElement('w:p')
        if ri < len(col_data) and col_data[ri]:
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = str(col_data[ri])
            r.append(t)
            p.append(r)
        tc.append(p)

        tr = row._tr
        existing_tcs = tr.findall(qn('w:tc'))
        if position == "start":
            if existing_tcs:
                existing_tcs[0].addprevious(tc)
            else:
                tr.append(tc)
        elif position == "end":
            tr.append(tc)
        else:
            col_idx = int(position)
            if col_idx < len(existing_tcs):
                existing_tcs[col_idx].addprevious(tc)
            else:
                tr.append(tc)

    return f"Column added to table at block {block_index}, position '{position}'"


@docx_tool()
def delete_table_column(doc, filename, block_index, col_index):
    """Delete a specific column from a table identified by block_index."""
    try:
        bi = resolve_table_block(doc, block_index)
    except ValueError as e:
        return str(e)
    table = bi.obj
    col_index = int(col_index)
    if col_index < 0 or col_index >= len(table.columns):
        return f"Invalid column index. Table has {len(table.columns)} columns."

    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        gcs = grid.findall(qn('w:gridCol'))
        if col_index < len(gcs):
            grid.remove(gcs[col_index])

    for row in table.rows:
        tcs = row._tr.findall(qn('w:tc'))
        if col_index < len(tcs):
            row._tr.remove(tcs[col_index])

    return f"Column {col_index} deleted from table at block {block_index}"


# ---------------------------------------------------------------------------
# Paragraph text editing (addressed by block_index)
# ---------------------------------------------------------------------------

@docx_tool()
def set_paragraph_text(doc, filename, new_text, block_index=None,
                       target_text=None, preserve_formatting=True):
    """Replace the text of an existing paragraph block.

    Find by block_index or target_text. When preserve_formatting is True,
    keeps the first run's formatting and applies it to the new text.
    """
    bi = find_block(doc, target_text=target_text, block_index=block_index)
    if bi is None:
        return "Target paragraph not found"
    if bi.type == "table":
        return f"Block {bi.index} is a table, not a paragraph"

    para = bi.obj
    if preserve_formatting and para.runs:
        from word_document_server.operations.formatting import _capture_run_format, _restore_run_format
        fmt = _capture_run_format(para.runs[0])
        for run in para.runs:
            run._element.getparent().remove(run._element)
        new_run = para.add_run(new_text)
        _restore_run_format(new_run, fmt)
    else:
        for run in para.runs:
            run._element.getparent().remove(run._element)
        para.add_run(new_text)

    return f"Block {bi.index} text updated in {filename}"
