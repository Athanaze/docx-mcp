"""
Tests for list/numbering functionality.

Critical for legal documents: lists must render properly on ANY document,
not just those with pre-existing numbering definitions.
"""
import os
import zipfile
import pytest
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from word_document_server.operations.content import add_list, add_paragraph, insert_content
from word_document_server.operations.numbering import ensure_list_definitions


class TestNumberingDefinitions:
    """Verify that numbering.xml is properly created with all required elements."""

    def test_creates_numbering_part_on_blank(self, blank_docx):
        doc = Document(blank_docx)
        bullet_id, number_id = ensure_list_definitions(doc)
        doc.save(blank_docx)

        assert bullet_id is not None
        assert number_id is not None
        assert bullet_id != number_id

        with zipfile.ZipFile(blank_docx) as z:
            assert 'word/numbering.xml' in z.namelist()
            numbering_xml = z.read('word/numbering.xml')

        root = etree.fromstring(numbering_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        abstract_nums = root.findall('.//w:abstractNum', ns)
        assert len(abstract_nums) >= 2

        nums = root.findall('.//w:num', ns)
        assert len(nums) >= 2

    def test_idempotent(self, blank_docx):
        """Calling ensure_list_definitions twice should not duplicate entries."""
        doc = Document(blank_docx)
        id1 = ensure_list_definitions(doc)
        id2 = ensure_list_definitions(doc)
        assert id1 == id2

    def test_respects_existing_definitions(self, blank_docx):
        """If a document already has numbering definitions, don't conflict."""
        doc = Document(blank_docx)
        ensure_list_definitions(doc)
        doc.save(blank_docx)

        doc2 = Document(blank_docx)
        bid, nid = ensure_list_definitions(doc2)
        assert bid is not None
        assert nid is not None


class TestBulletList:
    def test_bullet_items_have_numPr(self, blank_docx):
        add_list(filename=blank_docx,
                 items=["Item A", "Item B", "Item C"],
                 list_type="bullet")

        doc = Document(blank_docx)
        list_paras = [p for p in doc.paragraphs if p.text in ("Item A", "Item B", "Item C")]
        assert len(list_paras) == 3

        for p in list_paras:
            pPr = p._element.find(qn('w:pPr'))
            assert pPr is not None, f"Paragraph '{p.text}' has no pPr"
            numPr = pPr.find(qn('w:numPr'))
            assert numPr is not None, f"Paragraph '{p.text}' has no numPr"
            ilvl = numPr.find(qn('w:ilvl'))
            assert ilvl is not None

    def test_bullet_numbering_xml_has_bullet_format(self, blank_docx):
        add_list(filename=blank_docx, items=["X"], list_type="bullet")

        with zipfile.ZipFile(blank_docx) as z:
            root = etree.fromstring(z.read('word/numbering.xml'))

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        abstract_nums = root.findall('.//w:abstractNum', ns)
        has_bullet = False
        for an in abstract_nums:
            for lvl in an.findall('.//w:lvl', ns):
                fmt = lvl.find('w:numFmt', ns)
                if fmt is not None and fmt.get(qn('w:val')) == 'bullet':
                    has_bullet = True
                    break
        assert has_bullet, "No bullet-type abstract numbering found"


class TestNumberedList:
    def test_numbered_items_have_numPr(self, blank_docx):
        add_list(filename=blank_docx,
                 items=["First", "Second"],
                 list_type="number")

        doc = Document(blank_docx)
        list_paras = [p for p in doc.paragraphs if p.text in ("First", "Second")]
        assert len(list_paras) == 2

        for p in list_paras:
            numPr = p._element.find(qn('w:pPr')).find(qn('w:numPr'))
            assert numPr is not None

    def test_numbered_format_is_decimal(self, blank_docx):
        add_list(filename=blank_docx, items=["X"], list_type="number")

        with zipfile.ZipFile(blank_docx) as z:
            root = etree.fromstring(z.read('word/numbering.xml'))

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        has_decimal = False
        for an in root.findall('.//w:abstractNum', ns):
            for lvl in an.findall('.//w:lvl', ns):
                fmt = lvl.find('w:numFmt', ns)
                if fmt is not None and fmt.get(qn('w:val')) == 'decimal':
                    has_decimal = True
                    break
        assert has_decimal, "No decimal-type abstract numbering found"


class TestInsertList:
    def test_insert_list_near_text(self, sample_docx):
        result = insert_content(
            filename=sample_docx, content_type="list",
            items=["Point 1", "Point 2"], target_text="first paragraph",
            position="after", list_type="bullet")
        assert "2 items" in result

        doc = Document(sample_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "Point 1" in texts
        assert "Point 2" in texts

    def test_insert_list_by_index(self, sample_docx):
        insert_content(
            filename=sample_docx, content_type="list",
            items=["A"], target_block_index=0,
            position="after", list_type="number")
        doc = Document(sample_docx)
        assert any(p.text == "A" for p in doc.paragraphs)


class TestMultipleListsRestart:
    """Verify that separate add_list calls create independent numbered lists."""

    def test_two_numbered_lists_restart(self, blank_docx):
        add_paragraph(filename=blank_docx, text="First list:")
        add_list(filename=blank_docx, items=["A", "B", "C"], list_type="number")
        add_paragraph(filename=blank_docx, text="Second list:")
        add_list(filename=blank_docx, items=["X", "Y"], list_type="number")

        doc = Document(blank_docx)
        # Find numIds used by list paragraphs
        num_ids = set()
        for para in doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    nid_el = numPr.find(qn('w:numId'))
                    if nid_el is not None:
                        num_ids.add(nid_el.get(qn('w:val')))

        # Each list should have its own numId
        assert len(num_ids) == 2, f"Expected 2 unique numIds, got {num_ids}"

    def test_bullet_lists_independent(self, blank_docx):
        add_list(filename=blank_docx, items=["A", "B"], list_type="bullet")
        add_paragraph(filename=blank_docx, text="Gap")
        add_list(filename=blank_docx, items=["X", "Y"], list_type="bullet")

        doc = Document(blank_docx)
        num_ids = set()
        for para in doc.paragraphs:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    nid_el = numPr.find(qn('w:numId'))
                    if nid_el is not None:
                        num_ids.add(nid_el.get(qn('w:val')))

        assert len(num_ids) == 2


class TestRoundTrip:
    """Verify lists survive save/reload cycles."""

    def test_bullet_round_trip(self, blank_docx):
        add_list(filename=blank_docx,
                 items=["Alpha", "Beta", "Gamma"],
                 list_type="bullet")

        doc = Document(blank_docx)
        doc.save(blank_docx)

        doc2 = Document(blank_docx)
        list_paras = [p for p in doc2.paragraphs if p.text in ("Alpha", "Beta", "Gamma")]
        assert len(list_paras) == 3

        for p in list_paras:
            numPr = p._element.find(qn('w:pPr')).find(qn('w:numPr'))
            assert numPr is not None, f"numPr lost after round-trip for '{p.text}'"
