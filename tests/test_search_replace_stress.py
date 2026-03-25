"""
Debug stress test — verbose prints to find where it hangs.
"""
import sys, os, shutil, time
sys.path.insert(0, '/home/s/Office-Word-MCP-Server')
sys.stdout.reconfigure(line_buffering=True)  # Force flush on every print

print("[0] Script started", flush=True)

print("[1] Importing docx...", flush=True)
t0 = time.time()
from docx import Document
from docx.shared import Pt, RGBColor
print(f"[1] docx imported in {time.time()-t0:.1f}s", flush=True)

print("[2] Importing find_and_replace_text...", flush=True)
from word_document_server.utils.document_utils import find_and_replace_text
print("[2] Imported OK", flush=True)

FIXTURES = '/home/s/Office-Word-MCP-Server/tests/fixtures'

# ======== Part A: Load each real fixture ========
for fname in ['par-hlink-frags.docx', 'run-char-style.docx', 'comments-rich-para.docx']:
    path = os.path.join(FIXTURES, fname)
    print(f"\n[A] Loading {fname}...", flush=True)
    t0 = time.time()
    doc = Document(path)
    print(f"[A] Loaded in {time.time()-t0:.1f}s — {len(doc.paragraphs)} paras", flush=True)
    
    multi = 0
    for p in doc.paragraphs:
        if len(p.runs) > 1:
            multi += 1
    print(f"[A] Multi-run paragraphs: {multi}", flush=True)
    
    # Find ONE cross-run span
    print(f"[A] Scanning for cross-run span...", flush=True)
    span_found = None
    para_idx = None
    for i, p in enumerate(doc.paragraphs):
        runs = p.runs
        if len(runs) < 2:
            continue
        for j in range(len(runs)-1):
            t1 = runs[j].text or ''
            t2 = runs[j+1].text or ''
            if len(t1) >= 2 and len(t2) >= 2:
                span = t1[-3:] + t2[:3]
                if len(span.strip()) > 3 and span in p.text:
                    # Check it's truly cross-run
                    in_single = any(span in (r.text or '') for r in runs)
                    if not in_single:
                        span_found = span
                        para_idx = i
                        break
        if span_found:
            break
    
    if span_found:
        print(f"[A] Found cross-run span at para[{para_idx}]: {repr(span_found)}", flush=True)
        
        # Work on copy
        wp = os.path.join(FIXTURES, f'_wc_{fname}')
        shutil.copy2(path, wp)
        doc2 = Document(wp)
        
        print(f"[A] Running find_and_replace_text...", flush=True)
        t0 = time.time()
        cnt = find_and_replace_text(doc2, span_found, 'XFIX')
        print(f"[A] Done in {time.time()-t0:.3f}s — count={cnt}", flush=True)
        
        p2 = doc2.paragraphs[para_idx]
        if 'XFIX' in p2.text:
            print(f"  ✅ Cross-run replace worked", flush=True)
        else:
            print(f"  ❌ Cross-run replace FAILED — text: {p2.text[:80]}", flush=True)
        
        # Save & reload
        print(f"[A] Saving...", flush=True)
        doc2.save(wp)
        doc3 = Document(wp)
        if 'XFIX' in doc3.paragraphs[para_idx].text:
            print(f"  ✅ Persists after reload", flush=True)
        else:
            print(f"  ❌ NOT persisted", flush=True)
        os.remove(wp)
    else:
        print(f"[A] No cross-run span found in this file", flush=True)

# ======== Part B: Synthetic edge cases ========
print(f"\n[B] Creating synthetic document...", flush=True)
doc = Document()

# B1: bold in middle
p = doc.add_paragraph(); p.add_run('normal '); r=p.add_run('BOLD'); r.bold=True; p.add_run(' normal')
# B2: 5 colored runs
p2 = doc.add_paragraph()
for t,c in [('alpha',RGBColor(255,0,0)),(' beta',RGBColor(0,255,0)),(' gamma',RGBColor(0,0,255)),(' delta',RGBColor(128,0,0)),(' eps',RGBColor(0,128,0))]:
    r=p2.add_run(t); r.font.color.rgb=c
# B3: unicode
p3 = doc.add_paragraph(); p3.add_run('café '); r=p3.add_run('résumé'); r.bold=True; p3.add_run(' naïve')
print(f"[B] Synthetic doc created: {len(doc.paragraphs)} paras", flush=True)

print(f"[B1] Testing 'al BOLD n' across 3 runs...", flush=True)
cnt = find_and_replace_text(doc, 'al BOLD n', 'FIXED')
print(f"[B1] count={cnt}, text='{doc.paragraphs[0].text}'", flush=True)
if cnt == 1: print("  ✅ B1 passed", flush=True)
else: print("  ❌ B1 FAILED", flush=True)

print(f"[B2] Testing 'beta gamma delta' across colored runs...", flush=True)
cnt = find_and_replace_text(doc, 'beta gamma delta', 'MERGED')
print(f"[B2] count={cnt}, text='{doc.paragraphs[1].text}'", flush=True)
if cnt == 1: print("  ✅ B2 passed", flush=True)
else: print("  ❌ B2 FAILED", flush=True)

print(f"[B3] Testing unicode 'café résumé naïve'...", flush=True)
cnt = find_and_replace_text(doc, 'café résumé naïve', 'UNICODE')
print(f"[B3] count={cnt}, text='{doc.paragraphs[2].text}'", flush=True)
if cnt == 1: print("  ✅ B3 passed", flush=True)
else: print("  ❌ B3 FAILED", flush=True)

print("\n🏁 Script completed!", flush=True)
