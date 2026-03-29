"""Tests for PDF conversion via LibreOffice."""
import os
import shutil
import pytest
from docx import Document

from word_document_server.operations.pdf import convert_to_pdf


def _has_libreoffice():
    return shutil.which('libreoffice') is not None


@pytest.mark.skipif(not _has_libreoffice(), reason="LibreOffice not installed")
class TestConvertToPdf:

    def test_basic_conversion(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_paragraph("Hello PDF")
        doc.save(path)

        result = convert_to_pdf(path)
        assert "Error" not in result
        pdf_path = os.path.join(tmp_dir, 'test.pdf')
        assert os.path.exists(pdf_path)
        assert os.path.getsize(pdf_path) > 500

    def test_custom_output_filename(self, tmp_dir):
        path = os.path.join(tmp_dir, 'input.docx')
        out = os.path.join(tmp_dir, 'output.pdf')
        doc = Document()
        doc.add_paragraph("Custom output")
        doc.save(path)

        result = convert_to_pdf(path, out)
        assert os.path.exists(out)

    def test_successive_calls(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_paragraph("Successive")
        doc.save(path)

        out1 = os.path.join(tmp_dir, 'out1.pdf')
        out2 = os.path.join(tmp_dir, 'out2.pdf')
        convert_to_pdf(path, out1)
        convert_to_pdf(path, out2)
        assert os.path.exists(out1)
        assert os.path.exists(out2)
