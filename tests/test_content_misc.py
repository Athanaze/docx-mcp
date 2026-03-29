"""Tests for content helpers that had no dedicated coverage: XML dump, TOC, picture."""
import struct
import zlib
from pathlib import Path

from docx import Document
from lxml import etree

from word_document_server.operations.content import (
    get_document_xml,
    add_table_of_contents,
    add_picture,
)


def _write_minimal_png(path: Path) -> None:
    """1×1 RGB PNG without third-party image libs (stdlib only)."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        crc = zlib.crc32(body) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + body + struct.pack(">I", crc)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b"\x00" + b"\xff\x00\x00"
    data = sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")
    path.write_bytes(data)


class TestGetDocumentXml:
    def test_returns_document_xml(self, blank_docx):
        out = get_document_xml(blank_docx)
        assert "w:document" in out
        root = etree.fromstring(out.encode("utf-8"))
        assert root.tag.endswith("document")

    def test_missing_file(self, tmp_path):
        missing = str(tmp_path / "nope.docx")
        out = get_document_xml(missing)
        assert "does not exist" in out.lower() or "not exist" in out.lower()


class TestAddTableOfContents:
    def test_inserts_toc_field(self, blank_docx):
        r = add_table_of_contents(filename=blank_docx, title="Contents", max_level=2)
        assert "table of contents" in r.lower() or "contents" in r.lower()
        xml = get_document_xml(blank_docx)
        assert "TOC" in xml
        assert "fldSimple" in xml


class TestAddPicture:
    def test_adds_inline_shape(self, blank_docx, tmp_path):
        png = tmp_path / "one.png"
        _write_minimal_png(png)
        r = add_picture(filename=blank_docx, image_path=str(png))
        assert "picture" in r.lower() or "added" in r.lower()
        doc = Document(blank_docx)
        assert len(doc.inline_shapes) == 1

    def test_missing_image(self, blank_docx):
        r = add_picture(filename=blank_docx, image_path="/nonexistent/pic.png")
        assert "does not exist" in r.lower() or "not exist" in r.lower()
