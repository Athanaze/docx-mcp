"""Tests that insert operations preserve formatting."""
import os
from docx import Document
from docx.shared import Pt, RGBColor

from word_document_server.operations.content import insert_content


class TestInsertFormattingInheritance:

    def test_inserted_paragraph_inherits_font(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("Styled paragraph")
        run.font.name = "Courier New"
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        doc.save(path)

        insert_content(filename=path, content_type="paragraph",
                       text="Inserted text", target_block_index=0,
                       position="after")

        doc2 = Document(path)
        inserted = doc2.paragraphs[1]
        assert inserted.text == "Inserted text"
        r = inserted.runs[0]
        assert r.font.name == "Courier New"
        assert r.font.size == Pt(14)
        assert r.bold is True
        assert r.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)

    def test_explicit_style_overrides_target(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("Bold Courier text")
        run.font.name = "Courier New"
        run.bold = True
        doc.save(path)

        insert_content(filename=path, content_type="paragraph",
                       text="Normal text", target_block_index=0,
                       position="after", style="Normal")

        doc2 = Document(path)
        inserted = doc2.paragraphs[1]
        assert inserted.text == "Normal text"
        assert inserted.runs[0].font.name is None or inserted.runs[0].font.name != "Courier New"

    def test_insert_before_preserves_formatting(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        p = doc.add_paragraph("")
        run = p.add_run("Italic Georgia text")
        run.font.name = "Georgia"
        run.italic = True
        doc.save(path)

        insert_content(filename=path, content_type="paragraph",
                       text="Before text", target_block_index=0,
                       position="before")

        doc2 = Document(path)
        inserted = doc2.paragraphs[0]
        assert inserted.text == "Before text"
        r = inserted.runs[0]
        assert r.font.name == "Georgia"
        assert r.italic is True

    def test_paragraph_near_heading_gets_normal(self, tmp_dir):
        """Inserting a paragraph near a heading should NOT inherit the heading style."""
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_paragraph("Target heading", style='Heading 1')
        doc.save(path)

        insert_content(filename=path, content_type="paragraph",
                       text="Inserted paragraph", target_block_index=0,
                       position="after")

        doc2 = Document(path)
        inserted = doc2.paragraphs[1]
        assert inserted.text == "Inserted paragraph"
        assert inserted.style.name == "Normal"

    def test_normal_style_inherited(self, tmp_dir):
        """Inserting a paragraph near a Normal paragraph SHOULD inherit its style."""
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_paragraph("Target paragraph", style='Normal')
        doc.save(path)

        insert_content(filename=path, content_type="paragraph",
                       text="Inserted text", target_block_index=0,
                       position="after")

        doc2 = Document(path)
        inserted = doc2.paragraphs[1]
        assert inserted.text == "Inserted text"
        assert inserted.style.name == "Normal"
