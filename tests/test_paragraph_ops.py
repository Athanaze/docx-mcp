"""Tests for paragraph operations: set_paragraph_text, move_block."""
import os
import pytest
from docx import Document
from docx.shared import Pt

from word_document_server.operations.content import (
    set_paragraph_text,
    move_block,
)


@pytest.fixture
def para_docx(tmp_path):
    path = str(tmp_path / "paras.docx")
    doc = Document()
    p0 = doc.add_paragraph("First paragraph")
    for r in p0.runs:
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = "Arial"
    doc.add_paragraph("Second paragraph")
    doc.add_paragraph("Third paragraph")
    doc.add_paragraph("Fourth paragraph")
    doc.save(path)
    return path


class TestSetParagraphText:
    def test_by_index_preserves_formatting(self, para_docx):
        result = set_paragraph_text(filename=para_docx, new_text="Updated first",
                                    block_index=0)
        assert "updated" in result.lower()
        doc = Document(para_docx)
        assert doc.paragraphs[0].text == "Updated first"
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].runs[0].font.name == "Arial"
        assert doc.paragraphs[0].runs[0].font.size == Pt(14)

    def test_by_search_text(self, para_docx):
        result = set_paragraph_text(filename=para_docx, new_text="New second",
                                    target_text="Second paragraph")
        assert "updated" in result.lower()
        doc = Document(para_docx)
        assert doc.paragraphs[1].text == "New second"

    def test_no_preserve_formatting(self, para_docx):
        set_paragraph_text(filename=para_docx, new_text="Plain text",
                           block_index=0, preserve_formatting=False)
        doc = Document(para_docx)
        assert doc.paragraphs[0].text == "Plain text"
        assert doc.paragraphs[0].runs[0].bold is None

    def test_not_found(self, para_docx):
        result = set_paragraph_text(filename=para_docx, new_text="X",
                                    target_text="nonexistent text here")
        assert "not found" in result.lower()

    def test_other_paragraphs_unchanged(self, para_docx):
        set_paragraph_text(filename=para_docx, new_text="Changed",
                           block_index=1)
        doc = Document(para_docx)
        assert doc.paragraphs[0].text == "First paragraph"
        assert doc.paragraphs[1].text == "Changed"
        assert doc.paragraphs[2].text == "Third paragraph"


class TestMoveBlock:
    def test_move_down(self, para_docx):
        result = move_block(
            filename=para_docx,
            source_index=0,
            target_index=2,
            position="after",
        )
        assert "moved" in result.lower()
        doc = Document(para_docx)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["Second paragraph", "Third paragraph",
                         "First paragraph", "Fourth paragraph"]

    def test_move_up(self, para_docx):
        move_block(
            filename=para_docx,
            source_index=3,
            target_index=0,
            position="before",
        )
        doc = Document(para_docx)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["Fourth paragraph", "First paragraph",
                         "Second paragraph", "Third paragraph"]

    def test_move_before(self, para_docx):
        move_block(
            filename=para_docx,
            source_index=2,
            target_index=0,
            position="before",
        )
        doc = Document(para_docx)
        texts = [p.text for p in doc.paragraphs]
        assert texts == ["Third paragraph", "First paragraph",
                         "Second paragraph", "Fourth paragraph"]

    def test_same_index(self, para_docx):
        result = move_block(
            filename=para_docx,
            source_index=1,
            target_index=1,
        )
        assert "same" in result.lower()

    def test_invalid_source(self, para_docx):
        result = move_block(
            filename=para_docx,
            source_index=99,
            target_index=0,
        )
        assert "invalid" in result.lower()

    def test_invalid_target(self, para_docx):
        result = move_block(
            filename=para_docx,
            source_index=0,
            target_index=99,
        )
        assert "invalid" in result.lower()
