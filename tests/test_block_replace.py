"""Tests for replace_block with invisible characters, tables, and anchors."""
import os
from docx import Document

from word_document_server.operations.content import replace_block


class TestReplaceBlockBelowHeader:

    def test_heading_with_nonbreaking_space(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Intro p1")
        h2 = doc.add_heading("", level=2)
        h2.add_run("Related\u00a0Work")
        doc.add_paragraph("Old content A")
        doc.add_paragraph("Old content B")
        doc.add_heading("Conclusion", level=2)
        doc.add_paragraph("End")
        doc.save(path)

        result = replace_block(filename=path, header_text="Related Work",
                               new_paragraphs=["New A", "New B"])
        assert "not found" not in result.lower()

        doc2 = Document(path)
        texts = [p.text for p in doc2.paragraphs]
        assert "Old content A" not in texts
        assert "New A" in texts
        assert "New B" in texts
        assert "Conclusion" in texts

    def test_heading_with_zero_width_space(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        h = doc.add_heading("", level=1)
        h.add_run("Con\u200bclusion")
        doc.add_paragraph("Old conclusion")
        doc.add_heading("References", level=1)
        doc.add_paragraph("[1] Ref")
        doc.save(path)

        result = replace_block(filename=path, header_text="Conclusion",
                               new_paragraphs=["Final thoughts."])
        assert "not found" not in result.lower()

        doc2 = Document(path)
        texts = [p.text for p in doc2.paragraphs]
        assert "Old conclusion" not in texts
        assert "Final thoughts." in texts
        assert "References" in texts


class TestReplaceBlockBetweenAnchors:

    def test_anchors_with_invisible_chars(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Intro")
        h = doc.add_heading("", level=2)
        h.add_run("Related\u00a0Work")
        doc.add_paragraph("Related A")
        doc.add_paragraph("Related B")
        h2 = doc.add_heading("", level=2)
        h2.add_run("Con\u200bclusion")
        doc.add_paragraph("End")
        doc.save(path)

        result = replace_block(filename=path,
                               start_anchor="Related Work",
                               end_anchor="Conclusion",
                               new_paragraphs=["Replaced."])
        assert "not found" not in result.lower()

        doc2 = Document(path)
        texts = [p.text for p in doc2.paragraphs]
        assert "Related A" not in texts
        assert "Replaced." in texts

    def test_no_end_anchor_stops_at_heading(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading("Section A", level=1)
        doc.add_paragraph("Content A1")
        doc.add_paragraph("Content A2")
        doc.add_heading("Section B", level=1)
        doc.add_paragraph("Content B")
        doc.save(path)

        result = replace_block(filename=path,
                               start_anchor="Section A",
                               new_paragraphs=["New A."])
        assert "not found" not in result.lower()

        doc2 = Document(path)
        texts = [p.text for p in doc2.paragraphs]
        assert "Content A1" not in texts
        assert "New A." in texts
        assert "Content B" in texts

    def test_tables_between_anchors_removed(self, tmp_dir):
        path = os.path.join(tmp_dir, 'test.docx')
        doc = Document()
        doc.add_heading("Section A", level=1)
        doc.add_paragraph("Before table")
        doc.add_table(rows=2, cols=2, style='Table Grid')
        doc.add_paragraph("After table")
        doc.add_heading("Section B", level=1)
        doc.add_paragraph("Section B content")
        doc.save(path)

        result = replace_block(filename=path,
                               start_anchor="Section A",
                               end_anchor="Section B",
                               new_paragraphs=["Clean section."])
        assert "not found" not in result.lower()

        doc2 = Document(path)
        assert len(doc2.tables) == 0
        texts = [p.text for p in doc2.paragraphs]
        assert "Clean section." in texts
        assert "Section B content" in texts
