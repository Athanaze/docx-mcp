"""Tests for the unified block item abstraction (blocks.py)."""
import os

import pytest
from docx import Document

from word_document_server.operations.blocks import (
    BlockItem, get_block_items, resolve_block,
    resolve_paragraph_block, resolve_table_block,
    find_block, normalize_text, _classify_paragraph,
    word_count_from_blocks,
)


class TestWordCountFromBlocks:
    def test_includes_table_cell_text(self, sample_docx):
        doc = Document(sample_docx)
        walk = sum(len((p.text or "").split()) for p in doc.paragraphs)
        wc = word_count_from_blocks(doc)
        assert wc > walk

    def test_blank_document(self, blank_docx):
        doc = Document(blank_docx)
        assert word_count_from_blocks(doc) == 0


class TestGetBlockItems:
    def test_empty_document(self, tmp_dir):
        path = os.path.join(tmp_dir, "empty.docx")
        doc = Document()
        doc.save(path)
        doc = Document(path)
        items = get_block_items(doc)
        assert isinstance(items, list)

    def test_mixed_content(self, sample_docx):
        doc = Document(sample_docx)
        items = get_block_items(doc)
        types = [bi.type for bi in items]
        assert "heading" in types
        assert "table" in types
        assert "paragraph" in types

    def test_indices_are_sequential(self, sample_docx):
        doc = Document(sample_docx)
        items = get_block_items(doc)
        for i, bi in enumerate(items):
            assert bi.index == i

    def test_returns_block_item_namedtuples(self, sample_docx):
        doc = Document(sample_docx)
        items = get_block_items(doc)
        for bi in items:
            assert isinstance(bi, BlockItem)
            assert bi.type in ("paragraph", "heading", "list_item", "table")


class TestClassifyParagraph:
    def test_heading(self, tmp_dir):
        path = os.path.join(tmp_dir, "h.docx")
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.save(path)
        doc = Document(path)
        items = get_block_items(doc)
        heading_items = [bi for bi in items if bi.type == "heading"]
        assert len(heading_items) >= 1

    def test_list_item(self, tmp_dir):
        from word_document_server.operations.numbering import (
            ensure_list_definitions, set_paragraph_list,
        )
        path = os.path.join(tmp_dir, "list.docx")
        doc = Document()
        ensure_list_definitions(doc)
        p = doc.add_paragraph("Item 1")
        set_paragraph_list(p, 1, level=0)
        doc.save(path)
        doc = Document(path)
        items = get_block_items(doc)
        list_items = [bi for bi in items if bi.type == "list_item"]
        assert len(list_items) >= 1


class TestResolveBlock:
    def test_valid_index(self, sample_docx):
        doc = Document(sample_docx)
        bi = resolve_block(doc, 0)
        assert bi.index == 0

    def test_invalid_index_raises(self, sample_docx):
        doc = Document(sample_docx)
        with pytest.raises(ValueError, match="Invalid block index"):
            resolve_block(doc, 999)

    def test_negative_index_raises(self, sample_docx):
        doc = Document(sample_docx)
        with pytest.raises(ValueError, match="Invalid block index"):
            resolve_block(doc, -1)


class TestResolveParagraphBlock:
    def test_returns_paragraph(self, sample_docx):
        doc = Document(sample_docx)
        bi = resolve_paragraph_block(doc, 0)
        assert bi.type in ("paragraph", "heading", "list_item")

    def test_raises_for_table(self, sample_docx):
        doc = Document(sample_docx)
        items = get_block_items(doc)
        table_idx = next(bi.index for bi in items if bi.type == "table")
        with pytest.raises(ValueError, match="table"):
            resolve_paragraph_block(doc, table_idx)


class TestResolveTableBlock:
    def test_returns_table(self, sample_docx):
        doc = Document(sample_docx)
        items = get_block_items(doc)
        table_idx = next(bi.index for bi in items if bi.type == "table")
        bi = resolve_table_block(doc, table_idx)
        assert bi.type == "table"

    def test_raises_for_paragraph(self, sample_docx):
        doc = Document(sample_docx)
        with pytest.raises(ValueError, match="not a table"):
            resolve_table_block(doc, 0)


class TestNormalizeText:
    def test_strips_zero_width_chars(self):
        assert normalize_text("he\u200bllo") == "hello"

    def test_collapses_whitespace(self):
        assert normalize_text("hello   world") == "hello world"

    def test_nfc_normalization(self):
        assert normalize_text("e\u0301") == "\u00e9"

    def test_nbsp_to_space(self):
        assert normalize_text("hello\u00a0world") == "hello world"


class TestFindBlock:
    def test_find_by_index(self, sample_docx):
        doc = Document(sample_docx)
        bi = find_block(doc, block_index=0)
        assert bi is not None
        assert bi.index == 0

    def test_find_by_index_out_of_range(self, sample_docx):
        doc = Document(sample_docx)
        assert find_block(doc, block_index=999) is None

    def test_find_by_text_exact(self, sample_docx):
        doc = Document(sample_docx)
        bi = find_block(doc, target_text="Introduction")
        assert bi is not None
        assert bi.type == "heading"

    def test_find_by_text_fuzzy(self, sample_docx):
        doc = Document(sample_docx)
        bi = find_block(doc, target_text="first paragraph")
        assert bi is not None

    def test_find_empty_text_returns_none(self, sample_docx):
        doc = Document(sample_docx)
        assert find_block(doc, target_text="") is None

    def test_find_no_args_returns_none(self, sample_docx):
        doc = Document(sample_docx)
        assert find_block(doc) is None
