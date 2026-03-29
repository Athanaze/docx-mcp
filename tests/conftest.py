"""
Shared pytest fixtures for Word MCP Server tests.
"""
import os
import tempfile
import shutil

import pytest
from docx import Document


@pytest.fixture
def tmp_dir():
    """Create a temporary directory that is cleaned up after the test."""
    d = tempfile.mkdtemp(prefix="word_mcp_test_")
    yield d
    shutil.rmtree(d, ignore_errors=True)


@pytest.fixture
def tmp_docx(tmp_dir):
    """Return a path to a temporary .docx file (does not create it)."""
    return os.path.join(tmp_dir, "test.docx")


@pytest.fixture
def blank_docx(tmp_dir):
    """Create and return path to a blank .docx file."""
    path = os.path.join(tmp_dir, "blank.docx")
    Document().save(path)
    return path


@pytest.fixture
def sample_docx(tmp_dir):
    """Create a sample document with headings, paragraphs, and a table."""
    path = os.path.join(tmp_dir, "sample.docx")
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first paragraph.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_heading("Details", level=2)
    doc.add_paragraph("Some detail text here.")
    table = doc.add_table(rows=3, cols=3)
    for ri in range(3):
        for ci in range(3):
            table.cell(ri, ci).text = f"R{ri}C{ci}"
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph("Final paragraph.")
    doc.save(path)
    return path


@pytest.fixture(autouse=True)
def _clear_document_root(monkeypatch):
    """Ensure DOCUMENT_ROOT is not set during tests unless explicitly set."""
    monkeypatch.delenv('DOCUMENT_ROOT', raising=False)
