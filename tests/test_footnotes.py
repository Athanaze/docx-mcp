"""
Tests for footnote operations using robust OOXML manipulation.
"""
import json
import zipfile
import pytest
from lxml import etree
from docx import Document

from word_document_server.operations.content import add_paragraph, add_heading
from word_document_server.operations.footnotes import (
    add_footnote, delete_footnote, validate_footnotes,
)

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


class TestAddFootnote:
    def test_add_by_block_index(self, blank_docx):
        add_paragraph(filename=blank_docx, text="This needs a footnote.")
        result = add_footnote(blank_docx, block_index=0,
                              footnote_text="See appendix A.")
        data = json.loads(result)
        assert data["success"] is True
        assert data["footnote_id"] >= 2

    def test_add_by_search_text(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Important legal clause.")
        result = add_footnote(blank_docx, search_text="legal clause",
                              footnote_text="Reference: Case 123")
        data = json.loads(result)
        assert data["success"] is True

    def test_footnotes_xml_created(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Text with footnote.")
        add_footnote(blank_docx, block_index=0, footnote_text="FN text")

        with zipfile.ZipFile(blank_docx) as z:
            assert 'word/footnotes.xml' in z.namelist()
            root = etree.fromstring(z.read('word/footnotes.xml'))
            ns = {'w': W_NS}
            footnotes = root.findall('.//w:footnote', ns)
            user_fns = [f for f in footnotes
                        if f.get(f'{{{W_NS}}}id') not in ('-1', '0')]
            assert len(user_fns) >= 1

    def test_footnote_reference_in_document(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Ref test.")
        add_footnote(blank_docx, block_index=0, footnote_text="Ref content")

        with zipfile.ZipFile(blank_docx) as z:
            doc_root = etree.fromstring(z.read('word/document.xml'))
            ns = {'w': W_NS}
            refs = doc_root.xpath('//w:footnoteReference', namespaces=ns)
            assert len(refs) >= 1

    def test_position_before(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Before position test.")
        result = add_footnote(blank_docx, block_index=0,
                              footnote_text="Before fn", position="before")
        data = json.loads(result)
        assert data["success"] is True
        assert data["position"] == "before"

    def test_position_after(self, blank_docx):
        add_paragraph(filename=blank_docx, text="After position test.")
        result = add_footnote(blank_docx, block_index=0,
                              footnote_text="After fn", position="after")
        data = json.loads(result)
        assert data["position"] == "after"

    def test_multiple_footnotes(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Para one.")
        add_paragraph(filename=blank_docx, text="Para two.")
        r1 = json.loads(add_footnote(blank_docx, block_index=0,
                                     footnote_text="FN 1"))
        r2 = json.loads(add_footnote(blank_docx, block_index=1,
                                     footnote_text="FN 2"))
        assert r1["footnote_id"] != r2["footnote_id"]

    def test_text_not_found(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Some text.")
        result = add_footnote(blank_docx, search_text="nonexistent",
                              footnote_text="fn")
        assert "not found" in result.lower()

    def test_must_provide_locator(self, blank_docx):
        result = add_footnote(blank_docx, footnote_text="orphan")
        assert "must provide" in result.lower()


class TestDeleteFootnote:
    def test_delete_by_id(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Delete test.")
        r = json.loads(add_footnote(blank_docx, block_index=0,
                                    footnote_text="To be deleted"))
        fn_id = r["footnote_id"]

        result = delete_footnote(blank_docx, footnote_id=fn_id)
        data = json.loads(result)
        assert data["success"] is True

        with zipfile.ZipFile(blank_docx) as z:
            doc_root = etree.fromstring(z.read('word/document.xml'))
            ns = {'w': W_NS}
            refs = doc_root.xpath(
                f'//w:footnoteReference[@w:id="{fn_id}"]', namespaces=ns)
            assert len(refs) == 0

    def test_delete_by_search_text(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Searchable text for deletion.")
        add_footnote(blank_docx, search_text="Searchable text",
                     footnote_text="FN to delete")

        result = delete_footnote(blank_docx, search_text="Searchable text")
        data = json.loads(result)
        assert data["success"] is True

    def test_delete_nonexistent(self, blank_docx):
        add_paragraph(filename=blank_docx, text="No footnotes here.")
        result = delete_footnote(blank_docx, footnote_id=999)
        assert "not found" in result.lower() or "No footnotes" in result


class TestValidateFootnotes:
    def test_valid_document(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Validated doc.")
        add_footnote(blank_docx, block_index=0, footnote_text="Valid FN")

        result = validate_footnotes(blank_docx)
        data = json.loads(result)
        assert data["valid"] is True
        assert data["total_references"] >= 1
        assert data["total_content"] >= 1

    def test_empty_document(self, blank_docx):
        result = validate_footnotes(blank_docx)
        data = json.loads(result)
        assert data["total_references"] == 0


class TestFootnoteRoundTrip:
    """Verify footnotes survive save/reload cycles."""

    def test_round_trip(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Round trip test.")
        add_footnote(blank_docx, block_index=0, footnote_text="Persisted FN")

        doc = Document(blank_docx)
        doc.save(blank_docx)

        with zipfile.ZipFile(blank_docx) as z:
            assert 'word/footnotes.xml' in z.namelist()
            root = etree.fromstring(z.read('word/footnotes.xml'))
            ns = {'w': W_NS}
            user_fns = [f for f in root.findall('.//w:footnote', ns)
                        if f.get(f'{{{W_NS}}}id') not in ('-1', '0')]
            assert len(user_fns) >= 1


class TestFootnotePositioning:
    """Verify footnote reference is placed at the search text, not paragraph end."""

    def test_reference_after_search_text(self, blank_docx):
        add_paragraph(filename=blank_docx,
                      text="See Smith v. Jones for the ruling on damages.")
        add_footnote(blank_docx, search_text="Smith v. Jones",
                     footnote_text="Citation note", position="after")

        with zipfile.ZipFile(blank_docx) as z:
            doc_xml = z.read('word/document.xml')
            doc_root = etree.fromstring(doc_xml)
            ns = {'w': W_NS}
            para = doc_root.xpath('//w:p', namespaces=ns)[0]
            runs = list(para)

            ref_idx = None
            for i, child in enumerate(runs):
                if child.find(f'.//{{{W_NS}}}footnoteReference') is not None:
                    ref_idx = i
                    break
            assert ref_idx is not None

            # Collect text before and after footnote ref
            text_before = ""
            for child in runs[:ref_idx]:
                text_before += ''.join(
                    t.text or '' for t in child.iter(f'{{{W_NS}}}t'))
            text_after = ""
            for child in runs[ref_idx + 1:]:
                text_after += ''.join(
                    t.text or '' for t in child.iter(f'{{{W_NS}}}t'))

            assert "Smith v. Jones" in text_before
            assert "for the ruling" in text_after
