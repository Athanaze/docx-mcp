"""Tests for document preview (docx → PDF → PNG); optional system tools."""
import json
import os
import shutil
import struct
import zlib

import pytest
from docx import Document

from word_document_server.operations.preview import (
    compare_rendered_pages,
    preview_document,
    render_document_pages,
)


def _has_preview_stack():
    lo = shutil.which("libreoffice") or shutil.which("soffice")
    return bool(lo and shutil.which("pdftoppm"))


def _write_png(path, rgb):
    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        crc = zlib.crc32(body) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + body + struct.pack(">I", crc)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw = b"\x00" + bytes(rgb)
    data = sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", zlib.compress(raw)) + chunk(b"IEND", b"")
    path.write_bytes(data)


@pytest.mark.skipif(not _has_preview_stack(), reason="LibreOffice and/or pdftoppm not installed")
class TestPreviewDocument:
    def test_produces_png_and_pdf(self, tmp_path):
        docx_path = tmp_path / "prev.docx"
        d = Document()
        d.add_paragraph("Preview line one.")
        d.save(str(docx_path))

        out_dir = tmp_path / "out"
        out_dir.mkdir()
        result = preview_document(str(docx_path), output_dir=str(out_dir))
        data = json.loads(result)
        assert data.get("success") is True
        assert data.get("count", 0) >= 1
        for p in data["pages"]:
            assert os.path.isfile(p)
            assert p.endswith(".png")
        assert os.path.isfile(data["pdf"])
        assert len(data.get("page_images", [])) == data.get("count")
        first = data["page_images"][0]
        assert os.path.isfile(first["path"])
        assert first["size_bytes"] > 0
        assert first["page_number"] >= 1

    def test_render_alias(self, tmp_path):
        docx_path = tmp_path / "prev2.docx"
        d = Document()
        d.add_paragraph("Preview alias.")
        d.save(str(docx_path))

        out_dir = tmp_path / "out2"
        out_dir.mkdir()
        data = json.loads(render_document_pages(str(docx_path), output_dir=str(out_dir)))
        assert data.get("success") is True
        assert data.get("count", 0) >= 1


class TestCompareRenderedPages:
    def test_detects_changed_pages(self, tmp_path):
        b1 = tmp_path / "before_page-001.png"
        b2 = tmp_path / "before_page-002.png"
        a1 = tmp_path / "after_page-001.png"
        a2 = tmp_path / "after_page-002.png"
        _write_png(b1, (255, 0, 0))
        _write_png(b2, (0, 255, 0))
        _write_png(a1, (255, 0, 0))
        _write_png(a2, (0, 0, 255))

        res = json.loads(compare_rendered_pages(
            before_pages=[str(b1), str(b2)],
            after_pages=[str(a1), str(a2)],
            change_threshold_percent=0.1,
        ))
        assert res["success"] is True
        assert res["matched_count"] == 2
        assert res["changed_count"] == 1
        assert "2" in res["changed_pages"]
        matched2 = [p for p in res["per_page"] if p["page_key"] == "2"][0]
        assert matched2["pixel_diff_percent"] > 0

    def test_handles_added_removed_pages(self, tmp_path):
        b1 = tmp_path / "before_page-001.png"
        a1 = tmp_path / "after_page-001.png"
        a2 = tmp_path / "after_page-002.png"
        _write_png(b1, (255, 0, 0))
        _write_png(a1, (255, 0, 0))
        _write_png(a2, (0, 255, 0))

        res = json.loads(compare_rendered_pages(
            before_pages=[str(b1)],
            after_pages=[str(a1), str(a2)],
        ))
        assert res["success"] is True
        assert "2" in res["added_pages"]
        assert res["changed_count"] == 1
