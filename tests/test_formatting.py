"""
Tests for formatting operations: text formatting, styles, table formatting.
"""
import os
import json
import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from word_document_server.operations.content import (
    add_paragraph, add_table, add_heading,
)
from word_document_server.operations.formatting import (
    format_text, create_style, format_table,
    set_column_widths, merge_cells, merge_cells_horizontal,
    merge_cells_vertical, set_table_cell_shading,
    apply_table_alternating_rows, highlight_table_header,
    set_cell_alignment, format_cell_text, set_cell_padding,
    set_table_width, auto_fit_table, set_table_alignment_all,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


@pytest.fixture
def doc_with_text(blank_docx):
    """Document with a paragraph of known text."""
    add_paragraph(filename=blank_docx, text="Hello World Test")
    return blank_docx


@pytest.fixture
def doc_with_table(blank_docx):
    """Document with a 3x3 table."""
    add_table(filename=blank_docx, rows=3, cols=3,
              data=[["A", "B", "C"], ["D", "E", "F"], ["G", "H", "I"]])
    return blank_docx


class TestFormatText:
    def test_bold_range(self, doc_with_text):
        result = format_text(filename=doc_with_text, block_index=0,
                             start_pos=6, end_pos=11, bold=True)
        assert "World" in result
        doc = Document(doc_with_text)
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) >= 1
        assert any("World" in r.text for r in bold_runs)

    def test_color_range(self, doc_with_text):
        format_text(filename=doc_with_text, block_index=0,
                    start_pos=0, end_pos=5, color="FF0000")
        doc = Document(doc_with_text)
        run = doc.paragraphs[0].runs[0]
        assert run.font.color.rgb == RGBColor(0xFF, 0, 0)

    def test_invalid_paragraph(self, doc_with_text):
        result = format_text(filename=doc_with_text, block_index=99,
                             start_pos=0, end_pos=1)
        assert "invalid" in result.lower() or "Invalid" in result

    def test_invalid_positions(self, doc_with_text):
        result = format_text(filename=doc_with_text, block_index=0,
                             start_pos=100, end_pos=200)
        assert "invalid" in result.lower() or "Invalid" in result

    def test_search_text_bold(self, doc_with_text):
        result = format_text(filename=doc_with_text, search_text="World",
                             bold=True)
        assert "World" in result
        doc = Document(doc_with_text)
        bold_runs = [r for r in doc.paragraphs[0].runs if r.bold]
        assert any("World" in r.text for r in bold_runs)

    def test_search_text_not_found(self, doc_with_text):
        result = format_text(filename=doc_with_text, search_text="Nonexistent")
        assert "not found" in result.lower()

    def test_search_text_with_paragraph_scope(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Apple pie")
        add_paragraph(filename=blank_docx, text="Apple sauce")
        result = format_text(filename=blank_docx, search_text="Apple",
                             block_index=1, italic=True)
        assert "Apple" in result
        doc = Document(blank_docx)
        assert doc.paragraphs[0].runs[0].italic is not True
        italic_runs = [r for r in doc.paragraphs[1].runs if r.italic]
        assert any("Apple" in r.text for r in italic_runs)

    def test_search_text_occurrence(self, blank_docx):
        add_paragraph(filename=blank_docx, text="one two one three")
        result = format_text(filename=blank_docx, search_text="one",
                             match_occurrence=2, bold=True)
        assert "one" in result
        doc = Document(blank_docx)
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) == 1
        non_bold = [r for r in runs if not r.bold]
        assert any("one" in r.text for r in non_bold)


class TestCreateStyle:
    def test_creates_style(self, blank_docx):
        result = create_style(filename=blank_docx, style_name="Legal Body",
                              font_size=12, font_name="Times New Roman")
        assert "created" in result.lower()
        doc = Document(blank_docx)
        style = doc.styles["Legal Body"]
        assert style.font.size == Pt(12)
        assert style.font.name == "Times New Roman"

    def test_duplicate_style(self, blank_docx):
        create_style(filename=blank_docx, style_name="MyStyle")
        result = create_style(filename=blank_docx, style_name="MyStyle")
        assert "already exists" in result.lower()


class TestFormatTable:
    def test_header_row(self, doc_with_table):
        result = format_table(filename=doc_with_table, block_index=0,
                              has_header_row=True)
        assert "formatted" in result.lower()
        doc = Document(doc_with_table)
        first_row_cell = doc.tables[0].rows[0].cells[0]
        assert any(r.bold for r in first_row_cell.paragraphs[0].runs)

    def test_border_style(self, doc_with_table):
        result = format_table(filename=doc_with_table, block_index=0,
                              border_style="double")
        assert "formatted" in result.lower()

    def test_invalid_table(self, blank_docx):
        result = format_table(filename=blank_docx, block_index=0)
        assert "invalid" in result.lower() or "Invalid" in result


class TestColumnWidths:
    def test_set_widths(self, doc_with_table):
        result = set_column_widths(filename=doc_with_table, block_index=0,
                                   widths=[100, 150, 200])
        assert "set" in result.lower()


class TestMergeCells:
    def test_merge_rectangular(self, doc_with_table):
        result = merge_cells(filename=doc_with_table, block_index=0,
                             start_row=0, start_col=0, end_row=1, end_col=1)
        assert "merged" in result.lower()
        doc = Document(doc_with_table)
        # After merge, cell(0,0) and cell(1,1) share the same underlying tc element
        assert doc.tables[0].cell(0, 0)._tc is doc.tables[0].cell(1, 1)._tc

    def test_merge_horizontal(self, doc_with_table):
        result = merge_cells_horizontal(
            filename=doc_with_table, block_index=0,
            row=0, start_col=0, end_col=1,
        )
        assert "merged" in result.lower()
        doc = Document(doc_with_table)
        assert doc.tables[0].cell(0, 0)._tc is doc.tables[0].cell(0, 1)._tc

    def test_merge_vertical(self, doc_with_table):
        result = merge_cells_vertical(
            filename=doc_with_table, block_index=0,
            col=0, start_row=0, end_row=1,
        )
        assert "merged" in result.lower()
        doc = Document(doc_with_table)
        assert doc.tables[0].cell(0, 0)._tc is doc.tables[0].cell(1, 0)._tc


class TestTableWidthAndLayout:
    def test_set_table_width_points(self, doc_with_table):
        r = set_table_width(filename=doc_with_table, block_index=0,
                            width=400, width_type="points")
        assert "width" in r.lower()

    def test_set_table_width_auto(self, doc_with_table):
        r = set_table_width(filename=doc_with_table, block_index=0,
                            width=0, width_type="auto")
        assert "width" in r.lower()

    def test_auto_fit_table(self, doc_with_table):
        r = auto_fit_table(filename=doc_with_table, block_index=0)
        assert "auto" in r.lower() or "fit" in r.lower()

    def test_set_table_alignment_all(self, doc_with_table):
        r = set_table_alignment_all(
            filename=doc_with_table, block_index=0,
            horizontal="center", vertical="bottom",
        )
        assert "alignment" in r.lower()
        doc = Document(doc_with_table)
        cell = doc.tables[0].cell(0, 0)
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM


class TestCellShading:
    def test_shading(self, doc_with_table):
        result = set_table_cell_shading(filename=doc_with_table, block_index=0,
                                        row_index=0, col_index=0,
                                        fill_color="FF0000")
        assert "shading" in result.lower() or "Shading" in result


class TestAlternatingRows:
    def test_alternating(self, doc_with_table):
        result = apply_table_alternating_rows(filename=doc_with_table,
                                              block_index=0)
        assert "alternating" in result.lower()


class TestHighlightHeader:
    def test_highlight(self, doc_with_table):
        result = highlight_table_header(filename=doc_with_table, block_index=0)
        assert "highlighted" in result.lower() or "Header" in result


class TestCellAlignment:
    def test_alignment(self, doc_with_table):
        result = set_cell_alignment(filename=doc_with_table, block_index=0,
                                    row_index=0, col_index=0,
                                    horizontal="center", vertical="center")
        assert "alignment" in result.lower() or "Alignment" in result


class TestFormatCellText:
    def test_format_cell(self, doc_with_table):
        result = format_cell_text(filename=doc_with_table, block_index=0,
                                  row_index=0, col_index=0,
                                  text_content="New Text", bold=True)
        assert "formatted" in result.lower() or "Cell" in result
        doc = Document(doc_with_table)
        assert doc.tables[0].cell(0, 0).text == "New Text"


class TestCellPadding:
    def test_padding(self, doc_with_table):
        result = set_cell_padding(filename=doc_with_table, block_index=0,
                                  row_index=0, col_index=0,
                                  top=5.0, bottom=5.0)
        assert "padding" in result.lower() or "Padding" in result
