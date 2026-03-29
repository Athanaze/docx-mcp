"""
Regression tests for search_and_replace.

Bug: when replace_text contained find_text as a prefix (e.g. replacing
"Sample Document" with "Sample Document (verified)"), _replace_in_runs
searched from index 0 after each replace and matched again inside the new
text, looping forever and hanging the MCP server.

Related: empty replace_text must still advance search_pos (e.g. "xx" -> remove
"x" twice) or the same infinite loop occurs.

Tests here exercise single-run, multi-run, table cells, multiple hits per
paragraph, and empty replacement, plus a wall-clock bound on Unix (SIGALRM).
"""
from __future__ import annotations

import os
import signal
import time

import pytest
from docx import Document

from word_document_server.operations.content import (
    add_paragraph,
    add_table,
    search_and_replace,
)


class TestSubstringReplacementDoesNotReMatch:
    """New text starts with (or contains) find_text; must not loop on first match."""

    def test_single_run_prefix_replacement(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Sample Document intro")
        search_and_replace(
            filename=blank_docx,
            find_text="Sample Document",
            replace_text="Sample Document (verified)",
        )
        doc = Document(blank_docx)
        assert doc.paragraphs[0].text == "Sample Document (verified) intro"

    def test_cross_run_prefix_replacement(self, tmp_path):
        """Forces first_ri != last_ri path in _replace_in_runs."""
        path = str(tmp_path / "split_runs.docx")
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Sample ")
        p.add_run("Document")
        p.add_run(" end")
        doc.save(path)

        t0 = time.perf_counter()
        search_and_replace(
            filename=path,
            find_text="Sample Document",
            replace_text="Sample Document (verified)",
        )
        assert time.perf_counter() - t0 < 5.0

        doc2 = Document(path)
        assert doc2.paragraphs[0].text == "Sample Document (verified) end"

    def test_two_occurrences_same_paragraph(self, blank_docx):
        add_paragraph(
            filename=blank_docx,
            text="A Sample Document B Sample Document C",
        )
        search_and_replace(
            filename=blank_docx,
            find_text="Sample Document",
            replace_text="Sample Document (ok)",
        )
        doc = Document(blank_docx)
        assert doc.paragraphs[0].text == (
            "A Sample Document (ok) B Sample Document (ok) C"
        )

    def test_table_cell_prefix_replacement(self, blank_docx):
        add_table(
            filename=blank_docx,
            rows=1,
            cols=1,
            data=[["Sample Document"]],
        )
        t0 = time.perf_counter()
        search_and_replace(
            filename=blank_docx,
            find_text="Sample Document",
            replace_text="Sample Document (in table)",
        )
        assert time.perf_counter() - t0 < 5.0
        doc = Document(blank_docx)
        assert doc.tables[0].cell(0, 0).text == "Sample Document (in table)"


class TestEmptyReplacementAdvances:
    def test_remove_repeated_single_char(self, blank_docx):
        add_paragraph(filename=blank_docx, text="xx")
        search_and_replace(filename=blank_docx, find_text="x", replace_text="")
        doc = Document(blank_docx)
        assert doc.paragraphs[0].text == ""


class TestEmptyFindRejected:
    def test_empty_find_text(self, blank_docx):
        add_paragraph(filename=blank_docx, text="x")
        r = search_and_replace(filename=blank_docx, find_text="", replace_text="y")
        assert "empty" in r.lower()


@pytest.mark.skipif(not hasattr(signal, "SIGALRM"), reason="no SIGALRM on this platform")
class TestSearchReplaceWallClockBound:
    """Fails the suite if search_and_replace blocks (hang regression)."""

    def test_cross_run_prefix_finishes_within_two_seconds(self, tmp_path):
        path = str(tmp_path / "alarm.docx")
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Sample ")
        p.add_run("Document")
        doc.save(path)

        def _timeout(*_a):
            raise AssertionError("search_and_replace exceeded 2s wall clock")

        signal.signal(signal.SIGALRM, _timeout)
        signal.alarm(2)
        try:
            search_and_replace(
                filename=path,
                find_text="Sample Document",
                replace_text="Sample Document (MCP verified)",
            )
        finally:
            signal.alarm(0)
            signal.signal(signal.SIGALRM, signal.SIG_DFL)

        assert os.path.getsize(path) > 0
