"""Tests for table structure operations: delete, add/delete rows/columns."""
import os
import pytest
from docx import Document

from word_document_server.operations.content import (
    delete_block, add_table_row, delete_table_row,
    add_table_column, delete_table_column,
)


@pytest.fixture
def table_docx(tmp_path):
    path = str(tmp_path / "table.docx")
    doc = Document()
    doc.add_paragraph("Before table")
    t = doc.add_table(rows=3, cols=3)
    for ri in range(3):
        for ci in range(3):
            t.cell(ri, ci).text = f"R{ri}C{ci}"
    doc.add_paragraph("After table")
    doc.save(path)
    return path


class TestDeleteTable:
    def test_delete_only_table(self, table_docx):
        result = delete_block(filename=table_docx, block_index=1)
        assert "deleted" in result.lower()
        doc = Document(table_docx)
        assert len(doc.tables) == 0
        assert doc.paragraphs[0].text == "Before table"
        assert doc.paragraphs[1].text == "After table"

    def test_delete_preserves_other_tables(self, tmp_path):
        path = str(tmp_path / "multi.docx")
        doc = Document()
        t1 = doc.add_table(rows=1, cols=1)
        t1.cell(0, 0).text = "First"
        t2 = doc.add_table(rows=1, cols=1)
        t2.cell(0, 0).text = "Second"
        doc.save(path)
        delete_block(filename=path, block_index=0)
        doc2 = Document(path)
        assert len(doc2.tables) == 1
        assert doc2.tables[0].cell(0, 0).text == "Second"

    def test_invalid_index(self, table_docx):
        result = delete_block(filename=table_docx, block_index=5)
        assert "invalid" in result.lower()


class TestAddTableRow:
    def test_append_row(self, table_docx):
        result = add_table_row(filename=table_docx, block_index=1,
                               row_data=["A", "B", "C"])
        assert "added" in result.lower()
        doc = Document(table_docx)
        assert len(doc.tables[0].rows) == 4
        assert doc.tables[0].cell(3, 0).text == "A"
        assert doc.tables[0].cell(3, 2).text == "C"

    def test_prepend_row(self, table_docx):
        add_table_row(filename=table_docx, block_index=1,
                      row_data=["X", "Y", "Z"], position="start")
        doc = Document(table_docx)
        assert len(doc.tables[0].rows) == 4
        assert doc.tables[0].cell(0, 0).text == "X"
        assert doc.tables[0].cell(1, 0).text == "R0C0"

    def test_insert_at_index(self, table_docx):
        add_table_row(filename=table_docx, block_index=1,
                      row_data=["M1", "M2", "M3"], position="1")
        doc = Document(table_docx)
        assert len(doc.tables[0].rows) == 4
        assert doc.tables[0].cell(0, 0).text == "R0C0"
        assert doc.tables[0].cell(1, 0).text == "M1"
        assert doc.tables[0].cell(2, 0).text == "R1C0"

    def test_empty_row(self, table_docx):
        add_table_row(filename=table_docx, block_index=1)
        doc = Document(table_docx)
        assert len(doc.tables[0].rows) == 4
        assert doc.tables[0].cell(3, 0).text == ""


class TestDeleteTableRow:
    def test_delete_middle_row(self, table_docx):
        result = delete_table_row(filename=table_docx, block_index=1, row_index=1)
        assert "deleted" in result.lower()
        doc = Document(table_docx)
        assert len(doc.tables[0].rows) == 2
        assert doc.tables[0].cell(0, 0).text == "R0C0"
        assert doc.tables[0].cell(1, 0).text == "R2C0"

    def test_delete_first_row(self, table_docx):
        delete_table_row(filename=table_docx, block_index=1, row_index=0)
        doc = Document(table_docx)
        assert len(doc.tables[0].rows) == 2
        assert doc.tables[0].cell(0, 0).text == "R1C0"

    def test_invalid_row_index(self, table_docx):
        result = delete_table_row(filename=table_docx, block_index=1, row_index=10)
        assert "invalid" in result.lower()


class TestAddTableColumn:
    def test_append_column(self, table_docx):
        result = add_table_column(filename=table_docx, block_index=1,
                                  col_data=["D0", "D1", "D2"])
        assert "added" in result.lower()
        doc = Document(table_docx)
        assert len(doc.tables[0].columns) == 4
        assert doc.tables[0].cell(0, 3).text == "D0"
        assert doc.tables[0].cell(2, 3).text == "D2"

    def test_prepend_column(self, table_docx):
        add_table_column(filename=table_docx, block_index=1,
                         col_data=["P0", "P1", "P2"], position="start")
        doc = Document(table_docx)
        assert len(doc.tables[0].columns) == 4
        assert doc.tables[0].cell(0, 0).text == "P0"
        assert doc.tables[0].cell(0, 1).text == "R0C0"

    def test_insert_at_index(self, table_docx):
        add_table_column(filename=table_docx, block_index=1,
                         col_data=["M0", "M1", "M2"], position="1")
        doc = Document(table_docx)
        assert len(doc.tables[0].columns) == 4
        assert doc.tables[0].cell(0, 0).text == "R0C0"
        assert doc.tables[0].cell(0, 1).text == "M0"
        assert doc.tables[0].cell(0, 2).text == "R0C1"

    def test_empty_column(self, table_docx):
        add_table_column(filename=table_docx, block_index=1)
        doc = Document(table_docx)
        assert len(doc.tables[0].columns) == 4


class TestDeleteTableColumn:
    def test_delete_middle_column(self, table_docx):
        result = delete_table_column(filename=table_docx, block_index=1, col_index=1)
        assert "deleted" in result.lower()
        doc = Document(table_docx)
        assert len(doc.tables[0].columns) == 2
        assert doc.tables[0].cell(0, 0).text == "R0C0"
        assert doc.tables[0].cell(0, 1).text == "R0C2"

    def test_delete_first_column(self, table_docx):
        delete_table_column(filename=table_docx, block_index=1, col_index=0)
        doc = Document(table_docx)
        assert len(doc.tables[0].columns) == 2
        assert doc.tables[0].cell(0, 0).text == "R0C1"

    def test_invalid_col_index(self, table_docx):
        result = delete_table_column(filename=table_docx, block_index=1, col_index=10)
        assert "invalid" in result.lower()
