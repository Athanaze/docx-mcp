"""
Tests for comment operations using native python-docx 1.2.0 API.
"""
import json
import pytest
from docx import Document

from word_document_server.operations.content import add_paragraph
from word_document_server.operations.comments import get_comments, add_comment


class TestAddComment:
    def test_adds_comment(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Paragraph to comment on.")
        result = add_comment(filename=blank_docx, block_index=0,
                             text="Review this", author="Test User")
        assert "comment added" in result.lower()

    def test_invalid_block(self, blank_docx):
        result = add_comment(filename=blank_docx, block_index=99,
                             text="No such paragraph")
        assert "invalid" in result.lower() or "Invalid" in result

    def test_comment_survives_round_trip(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Some text.")
        add_comment(filename=blank_docx, block_index=0,
                    text="My comment", author="Author A")

        doc = Document(blank_docx)
        doc.save(blank_docx)

        doc2 = Document(blank_docx)
        assert len(list(doc2.comments)) >= 1


class TestGetComments:
    def test_gets_comments(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Text here.")
        add_comment(filename=blank_docx, block_index=0,
                    text="Note 1", author="Alice")

        result = get_comments(filename=blank_docx)
        data = json.loads(result)
        assert data["count"] >= 1
        assert data["comments"][0]["author"] == "Alice"
        assert data["comments"][0]["text"] == "Note 1"

    def test_filter_by_author(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Text.")
        add_comment(filename=blank_docx, block_index=0,
                    text="C1", author="Alice")
        add_comment(filename=blank_docx, block_index=0,
                    text="C2", author="Bob")

        result = get_comments(filename=blank_docx, author="Bob")
        data = json.loads(result)
        assert data["count"] == 1
        assert data["comments"][0]["author"] == "Bob"

    def test_no_comments(self, blank_docx):
        result = get_comments(filename=blank_docx)
        data = json.loads(result)
        assert data["count"] == 0
