"""
Tests for content operations: create, add, delete, search/replace, block replace.
"""
import os
import json
import pytest
from docx import Document

from word_document_server.operations.content import (
    create_document, get_document_info, get_document_text,
    get_document_outline, get_blocks, list_document_styles,
    set_paragraph_style,
    list_documents, copy_document,
    add_heading, add_paragraph, add_table, add_page_break,
    add_list, insert_content, delete_block,
    search_and_replace, find_text, replace_block,
)


class TestCreateDocument:
    def test_creates_file(self, tmp_docx):
        result = create_document(filename=tmp_docx, title="Test", author="Bot")
        assert os.path.exists(tmp_docx)
        assert "created" in result.lower()
        doc = Document(tmp_docx)
        assert doc.core_properties.title == "Test"
        assert doc.core_properties.author == "Bot"

    def test_auto_extension(self, tmp_dir):
        path = os.path.join(tmp_dir, "nodot")
        create_document(filename=path)
        assert os.path.exists(path + ".docx")


class TestDocumentInfo:
    def test_returns_json(self, sample_docx):
        result = get_document_info(filename=sample_docx)
        info = json.loads(result)
        assert info["paragraph_count"] > 0
        assert info["table_count"] == 1
        assert "word_count" in info
        assert "word_count_paragraph_walk" in info
        # Unified count includes table cell text; paragraph walk does not.
        assert info["word_count"] >= info["word_count_paragraph_walk"]

    def test_nonexistent(self, tmp_dir):
        result = get_document_info(filename=os.path.join(tmp_dir, "nope.docx"))
        assert "does not exist" in result


class TestGetText:
    def test_extracts_paragraphs(self, sample_docx):
        text = get_document_text(filename=sample_docx)
        assert "Introduction" in text
        assert "first paragraph" in text

    def test_includes_table_text(self, sample_docx):
        text = get_document_text(filename=sample_docx)
        assert "R0C0" in text

    def test_table_row_and_cell_limits(self, sample_docx):
        text = get_document_text(
            filename=sample_docx,
            include_indices=True,
            max_table_rows=1,
            max_cells_per_row=2,
        )
        assert "more rows omitted" in text
        assert "(+1 cells)" in text
        assert "Row 0:" in text
        assert "Row 1:" not in text

    def test_table_cell_char_limit(self, sample_docx):
        text = get_document_text(
            filename=sample_docx,
            include_indices=True,
            max_chars_per_cell=3,
        )
        assert "R0C…" in text or "…" in text


class TestGetOutline:
    def test_returns_structure(self, sample_docx):
        result = get_document_outline(filename=sample_docx)
        data = json.loads(result)
        blocks = data["blocks"]
        assert len(blocks) > 0
        assert data.get("total_block_count") == len(blocks)
        assert data.get("returned_blocks") == len(blocks)
        tables = [b for b in blocks if b["type"] == "table"]
        assert len(tables) == 1
        assert tables[0]["rows"] == 3

    def test_max_blocks_truncates(self, sample_docx):
        data = json.loads(get_document_outline(filename=sample_docx, max_blocks=3))
        assert len(data["blocks"]) == 3
        assert data["truncated"] is True
        assert data["total_block_count"] > 3


class TestGetBlocks:
    def test_table_and_runs(self, sample_docx):
        data = json.loads(get_blocks(filename=sample_docx))
        assert data["block_count"] > 0
        tbl = next(b for b in data["blocks"] if b["type"] == "table")
        assert tbl["rows"] == 3
        assert len(tbl["cell_texts"]) == 3

    def test_range(self, sample_docx):
        data = json.loads(
            get_blocks(filename=sample_docx, start_block_index=0, end_block_index=0)
        )
        assert len(data["blocks"]) == 1


class TestListStyles:
    def test_lists_paragraph_styles(self, blank_docx):
        data = json.loads(list_document_styles(filename=blank_docx))
        names = {s["name"] for s in data["styles"]}
        assert "Normal" in names


class TestSetParagraphStyle:
    def test_applies_heading_style(self, blank_docx):
        add_paragraph(filename=blank_docx, text="X")
        set_paragraph_style(filename=blank_docx, block_index=0, style_name="Heading 1")
        doc = Document(blank_docx)
        assert doc.paragraphs[0].style.name.startswith("Heading")


class TestListDocuments:
    def test_finds_docx(self, sample_docx, tmp_dir):
        result = list_documents(tmp_dir)
        data = json.loads(result)
        assert any("sample.docx" in d for d in data["documents"])


class TestCopyDocument:
    def test_copies_file(self, sample_docx, tmp_dir):
        dest = os.path.join(tmp_dir, "copy.docx")
        result = copy_document(sample_docx, dest)
        assert os.path.exists(dest)
        assert "copied" in result.lower() or dest in result


class TestAddHeading:
    def test_adds_heading(self, blank_docx):
        add_heading(filename=blank_docx, text="My Heading", level=2)
        doc = Document(blank_docx)
        assert any("My Heading" in p.text for p in doc.paragraphs)

    def test_heading_with_border(self, blank_docx):
        add_heading(filename=blank_docx, text="Bordered", level=1,
                    border_bottom=True)
        doc = Document(blank_docx)
        para = [p for p in doc.paragraphs if "Bordered" in p.text][0]
        from docx.oxml.ns import qn
        pBdr = para._element.find(qn('w:pPr'))
        assert pBdr is not None


class TestAddParagraph:
    def test_adds_text(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Hello World")
        doc = Document(blank_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "Hello World" in texts

    def test_with_formatting(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Bold text",
                      bold=True, font_size=14, font_name="Arial")
        doc = Document(blank_docx)
        para = [p for p in doc.paragraphs if "Bold text" in p.text][0]
        assert para.runs[0].bold is True
        assert para.runs[0].font.name == "Arial"

    def test_with_color(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Red text", color="FF0000")
        doc = Document(blank_docx)
        para = [p for p in doc.paragraphs if "Red text" in p.text][0]
        from docx.shared import RGBColor
        assert para.runs[0].font.color.rgb == RGBColor(0xFF, 0x00, 0x00)


class TestAddTable:
    def test_creates_table(self, blank_docx):
        add_table(filename=blank_docx, rows=2, cols=3,
                  data=[["a", "b", "c"], ["d", "e", "f"]])
        doc = Document(blank_docx)
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "a"
        assert doc.tables[0].cell(1, 2).text == "f"


class TestAddPageBreak:
    def test_adds_break(self, blank_docx):
        add_paragraph(filename=blank_docx, text="Before break")
        add_page_break(filename=blank_docx)
        doc = Document(blank_docx)
        assert len(doc.paragraphs) >= 2


class TestAddList:
    def test_bullet_list(self, blank_docx):
        result = add_list(filename=blank_docx,
                          items=["Item 1", "Item 2", "Item 3"],
                          list_type="bullet")
        assert "3 items" in result
        doc = Document(blank_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "Item 1" in texts
        assert "Item 2" in texts

    def test_numbered_list(self, blank_docx):
        result = add_list(filename=blank_docx,
                          items=["First", "Second"],
                          list_type="number")
        assert "2 items" in result

    def test_list_creates_numbering_xml(self, blank_docx):
        """Verify that numbering.xml is properly created for blank documents."""
        add_list(filename=blank_docx,
                 items=["A", "B"],
                 list_type="bullet")
        import zipfile
        with zipfile.ZipFile(blank_docx) as z:
            assert 'word/numbering.xml' in z.namelist()

    def test_list_has_numPr(self, blank_docx):
        """Verify paragraphs have w:numPr elements."""
        add_list(filename=blank_docx,
                 items=["X", "Y"],
                 list_type="number")
        doc = Document(blank_docx)
        from docx.oxml.ns import qn
        list_paras = [p for p in doc.paragraphs if p.text in ("X", "Y")]
        for p in list_paras:
            numPr = p._element.find(qn('w:pPr'))
            assert numPr is not None
            numPr_el = numPr.find(qn('w:numPr'))
            assert numPr_el is not None


class TestInsertContent:
    def test_insert_paragraph_after(self, sample_docx):
        result = insert_content(
            filename=sample_docx, content_type="paragraph",
            text="Inserted after", target_text="first paragraph",
            position="after")
        assert "inserted" in result.lower()
        doc = Document(sample_docx)
        texts = [p.text for p in doc.paragraphs]
        idx_target = texts.index("This is the first paragraph.")
        assert texts[idx_target + 1] == "Inserted after"

    def test_insert_paragraph_before(self, sample_docx):
        insert_content(
            filename=sample_docx, content_type="paragraph",
            text="Before intro", target_text="Introduction",
            position="before")
        doc = Document(sample_docx)
        texts = [p.text for p in doc.paragraphs]
        idx = texts.index("Introduction")
        assert texts[idx - 1] == "Before intro"

    def test_insert_heading(self, sample_docx):
        insert_content(
            filename=sample_docx, content_type="heading",
            text="New Section", target_text="Conclusion",
            position="before", level=2)
        doc = Document(sample_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "New Section" in texts

    def test_insert_list(self, sample_docx):
        result = insert_content(
            filename=sample_docx, content_type="list",
            items=["Apple", "Banana"], target_text="first paragraph",
            position="after", list_type="bullet")
        assert "2 items" in result

    def test_insert_by_index(self, sample_docx):
        insert_content(
            filename=sample_docx, content_type="paragraph",
            text="By index", target_block_index=0, position="after")
        doc = Document(sample_docx)
        assert doc.paragraphs[1].text == "By index"


class TestDeleteBlock:
    def test_deletes(self, sample_docx):
        doc = Document(sample_docx)
        original_count = len(doc.paragraphs)
        delete_block(filename=sample_docx, block_index=1)
        doc = Document(sample_docx)
        assert len(doc.paragraphs) == original_count - 1

    def test_invalid_index(self, sample_docx):
        result = delete_block(filename=sample_docx, block_index=999)
        assert "invalid" in result.lower() or "Invalid" in result


class TestSearchAndReplace:
    def test_replaces_text(self, sample_docx):
        result = search_and_replace(
            filename=sample_docx, find_text="first", replace_text="1st")
        assert "1" in result  # at least 1 replacement
        doc = Document(sample_docx)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "1st paragraph" in text
        assert "first paragraph" not in text

    def test_not_found(self, sample_docx):
        result = search_and_replace(
            filename=sample_docx, find_text="nonexistent", replace_text="x")
        assert "not found" in result.lower()


class TestFindText:
    def test_finds_matches(self, sample_docx):
        result = find_text(filename=sample_docx, text_to_find="paragraph")
        data = json.loads(result)
        assert data["count"] >= 3
        assert data["returned"] == len(data["matches"])
        assert all("block_index" in m for m in data["matches"])

    def test_max_results_truncates_matches_not_count(self, sample_docx):
        data = json.loads(
            find_text(filename=sample_docx, text_to_find="paragraph", max_results=1)
        )
        assert data["count"] >= 3
        assert len(data["matches"]) == 1
        assert data["returned"] == 1
        assert data.get("truncated") is True

    def test_case_insensitive(self, sample_docx):
        result = find_text(filename=sample_docx, text_to_find="INTRODUCTION",
                           match_case=False)
        data = json.loads(result)
        assert data["count"] >= 1


class TestReplaceBlock:
    def test_replace_below_header(self, sample_docx):
        result = replace_block(
            filename=sample_docx, header_text="Details",
            new_paragraphs=["New detail 1", "New detail 2"])
        assert "Replaced" in result
        doc = Document(sample_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "New detail 1" in texts
        assert "Some detail text here." not in texts

    def test_replace_between_anchors(self, sample_docx):
        result = replace_block(
            filename=sample_docx, start_anchor="Introduction",
            end_anchor="Details",
            new_paragraphs=["Replaced content"])
        assert "Replaced" in result
        doc = Document(sample_docx)
        texts = [p.text for p in doc.paragraphs]
        assert "Replaced content" in texts
        assert "This is the first paragraph." not in texts

    def test_header_not_found(self, sample_docx):
        result = replace_block(
            filename=sample_docx, header_text="Nonexistent Header",
            new_paragraphs=["x"])
        assert "not found" in result.lower()


class TestGetDocumentTextIndices:
    def test_paragraph_indices(self, sample_docx):
        text = get_document_text(filename=sample_docx, include_indices=True)
        assert "[0]" in text
        assert "[1]" in text
        assert "(Heading" in text

    def test_table_inline(self, tmp_dir):
        path = os.path.join(tmp_dir, "tbl.docx")
        doc = Document()
        doc.add_paragraph("Before")
        doc.add_table(rows=2, cols=2).cell(0, 0).text = "X"
        doc.add_paragraph("After")
        doc.save(path)
        text = get_document_text(filename=path, include_indices=True)
        lines = text.split("\n")
        before_idx = next(i for i, l in enumerate(lines) if "Before" in l)
        table_idx = next(i for i, l in enumerate(lines) if "(Table 2x2)" in l)
        after_idx = next(i for i, l in enumerate(lines) if "After" in l)
        assert before_idx < table_idx < after_idx

    def test_plain_text_mode(self, sample_docx):
        text = get_document_text(filename=sample_docx, include_indices=False)
        assert "[0]" not in text
        assert "(Heading" not in text


class TestFindParagraphFuzzyMatch:
    """Regression: _find_paragraph must not match empty paragraphs via fuzzy search."""

    def test_empty_paragraph_not_matched(self, tmp_dir):
        path = os.path.join(tmp_dir, "fuzzy.docx")
        doc = Document()
        doc.add_paragraph("Header")
        doc.add_paragraph("")
        doc.add_paragraph("Target text to find")
        doc.save(path)
        result = insert_content(
            filename=path, content_type="paragraph", text="Inserted",
            target_text="Target text to find", position="after")
        doc2 = Document(path)
        texts = [p.text for p in doc2.paragraphs]
        target_idx = texts.index("Target text to find")
        assert texts[target_idx + 1] == "Inserted"

    def test_fuzzy_substring_match(self, tmp_dir):
        path = os.path.join(tmp_dir, "fuzzy2.docx")
        doc = Document()
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("Le récapitulatif des prétentions est le suivant :")
        doc.save(path)
        result = insert_content(
            filename=path, content_type="paragraph", text="After recap",
            target_text="Le récapitulatif des prétentions", position="after")
        doc2 = Document(path)
        texts = [p.text for p in doc2.paragraphs]
        recap_idx = next(i for i, t in enumerate(texts) if "récapitulatif" in t)
        assert texts[recap_idx + 1] == "After recap"


class TestInsertTable:
    def test_insert_table_after_paragraph(self, tmp_dir):
        path = os.path.join(tmp_dir, "ins_tbl.docx")
        doc = Document()
        doc.add_paragraph("Before")
        doc.add_paragraph("After")
        doc.save(path)
        result = insert_content(
            filename=path, content_type="table",
            target_block_index=0, position="after",
            table_rows=2, table_cols=2,
            table_data=[["H1", "H2"], ["V1", "V2"]])
        assert "Table" in result
        doc2 = Document(path)
        assert len(doc2.tables) == 1
        body = list(doc2.element.body)
        para_els = [doc2.paragraphs[0]._element, doc2.paragraphs[1]._element]
        tbl_el = doc2.tables[0]._tbl
        assert body.index(para_els[0]) < body.index(tbl_el) < body.index(para_els[1])


class TestMergeDocumentsOrder:
    def test_preserves_table_position(self, tmp_dir):
        src = os.path.join(tmp_dir, "src.docx")
        doc = Document()
        doc.add_paragraph("Before table")
        doc.add_table(rows=1, cols=1).cell(0, 0).text = "T"
        doc.add_paragraph("After table")
        doc.save(src)

        from word_document_server.operations.content import merge_documents
        tgt = os.path.join(tmp_dir, "merged.docx")
        merge_documents(target_filename=tgt, source_filenames=[src])
        text = get_document_text(filename=tgt, include_indices=True)
        lines = text.split("\n")
        before_idx = next(i for i, l in enumerate(lines) if "Before table" in l)
        table_idx = next(i for i, l in enumerate(lines) if "(Table 1x1)" in l)
        after_idx = next(i for i, l in enumerate(lines) if "After table" in l)
        assert before_idx < table_idx < after_idx
